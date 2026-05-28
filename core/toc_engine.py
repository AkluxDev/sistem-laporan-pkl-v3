"""
toc_engine.py
=============
Native Word TOC engine menggunakan field XML lengkap.

PERBAIKAN:
- TOC field menggunakan fldChar begin/separate/end yang benar
- w:dirty="true" pada semua fldChar
- instrText menggunakan TOC \\o "1-4" \\h \\z \\u
- Tidak ada TOC manual / fake TOC
- Placeholder paragraph benar agar Word tahu field ada

CATATAN TEKNIS:
- python-docx tidak punya API langsung untuk TOC
- Harus inject XML manual ke paragraph
- w:dirty="true" pada fldChar wajib agar Word auto-recalculate
- updateFields di settings.xml harus true (lihat docx_utils.py)
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from .docx_utils import set_run_fonts
from .paragraph_engine import enable_flow_control
from .styles import FONT_NAME, HEADING_SIZE
import copy


# ============================================================
# FUNGSI UTAMA: Buat paragraf judul "DAFTAR ISI"
# ============================================================

def add_toc_title(doc, title="DAFTAR ISI", style_name="Heading 1"):
    """
    Tambah judul 'DAFTAR ISI' sebagai heading native.
    
    PENTING: Heading 1 digunakan agar muncul di TOC sebagai entry pertama.
    Jika tidak ingin muncul di TOC, gunakan style khusus non-heading.
    """
    para = doc.add_paragraph()
    # Gunakan style 'TOC Heading' jika ada, fallback ke normal centered
    try:
        para.style = doc.styles['TOC Heading']
    except KeyError:
        para.style = doc.styles['Normal']
    
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = para.add_run(title)
    run.bold = True
    run.font.size = _pt_to_emu(12)
    
    # Tambah spacing setelah judul
    from docx_utils import Pt
    para.paragraph_format.space_after = Pt(6)
    
    return para


def _pt_to_emu(pt):
    """Convert points ke EMU untuk font size."""
    from docx_utils import Pt
    return Pt(pt)


# ============================================================
# FUNGSI UTAMA: Insert TOC Field Native Word
# ============================================================

def insert_toc_field(doc):
    """
    Insert TOC menggunakan native Word field XML.
    
    Struktur XML yang dihasilkan:
    
    <w:p>
      <w:pPr>
        <w:pStyle w:val="Normal"/>
      </w:pPr>
      <w:r>
        <w:fldChar w:fldCharType="begin" w:dirty="true"/>
      </w:r>
      <w:r>
        <w:instrText xml:space="preserve"> TOC \\o "1-4" \\h \\z \\u </w:instrText>
      </w:r>
      <w:r>
        <w:fldChar w:fldCharType="separate"/>
      </w:r>
      <w:r>
        <w:t>Klik kanan → Perbarui Field untuk memperbarui daftar isi.</w:t>
      </w:r>
      <w:r>
        <w:fldChar w:fldCharType="end" w:dirty="true"/>
      </w:r>
    </w:p>
    
    ALASAN w:dirty="true":
    - Memberitahu Word bahwa field harus dihitung ulang
    - Kombinasi dengan updateFields di settings.xml
    - Word akan auto-update TOC saat dokumen dibuka
    
    ALASAN instrText "TOC \\o "1-4" \\h \\z \\u":
    - \\o "1-4" = tampilkan Heading 1 sampai Heading 4
    - \\h = buat hyperlink di setiap entry
    - \\z = sembunyikan nomor halaman saat web layout
    - \\u = gunakan TC fields juga (opsional, untuk compatibility)
    """
    
    # Buat paragraf baru untuk menampung TOC field
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    
    # Bersihkan runs yang ada (jika paragraph baru sudah punya run default)
    for run in para.runs:
        run._element.getparent().remove(run._element)
    
    p_elem = para._element
    
    # ---- Run 1: fldChar begin ----
    r1 = OxmlElement('w:r')
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    fldChar_begin.set(qn('w:dirty'), 'true')  # WAJIB: dirty=true
    r1.append(fldChar_begin)
    p_elem.append(r1)
    
    # ---- Run 2: instrText ----
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = ' TOC \\o "1-4" \\h \\z \\u '
    r2.append(instrText)
    p_elem.append(r2)
    
    # ---- Run 3: fldChar separate ----
    r3 = OxmlElement('w:r')
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    fldChar_sep.set(qn('w:dirty'), 'true')  # WAJIB: dirty=true
    r3.append(fldChar_sep)
    p_elem.append(r3)
    
    # ---- Run 4: Placeholder text (ditampilkan sebelum Word update) ----
    r4 = OxmlElement('w:r')
    placeholder = OxmlElement('w:t')
    placeholder.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    placeholder.text = 'Daftar isi akan diperbarui secara otomatis saat dokumen dibuka.'
    r4.append(placeholder)
    p_elem.append(r4)
    
    # ---- Run 5: fldChar end ----
    r5 = OxmlElement('w:r')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    fldChar_end.set(qn('w:dirty'), 'true')  # WAJIB: dirty=true
    r5.append(fldChar_end)
    p_elem.append(r5)
    
    return para


# ============================================================
# FUNGSI HELPER: Tambah TOC lengkap (judul + field)
# ============================================================

def add_full_toc(doc, title="DAFTAR ISI"):
    """
    Shortcut: tambah judul DAFTAR ISI + TOC field sekaligus.
    
    Gunakan fungsi ini di doc_generator.py:
    
        from toc_engine import add_full_toc
        add_full_toc(doc, "DAFTAR ISI")
    
    Returns:
        tuple: (title_paragraph, toc_field_paragraph)
    """
    title_para = add_toc_title(doc, title)
    toc_para = insert_toc_field(doc)
    return title_para, toc_para


# ============================================================
# FUNGSI: Verifikasi TOC field ada di document
# ============================================================

def verify_toc_field_exists(doc):
    """
    Debug helper: cek apakah TOC field sudah ada di document.
    
    Returns:
        bool: True jika TOC field ditemukan
    """
    for para in doc.paragraphs:
        for child in para._element:
            if child.tag == qn('w:r'):
                for sub in child:
                    if sub.tag == qn('w:instrText'):
                        if sub.text and 'TOC' in sub.text:
                            return True
    return False


class TocEngine:
    """
    TocEngine V2 — Sistem Daftar Isi Profesional.
    Menggunakan native Word fields dengan struktur XML yang kompleks
    agar terdeteksi penuh sebagai field TOC oleh Microsoft Word.
    """
    def __init__(self, document):
        self.document = document
        self._placeholder_tag = "__PKL_TOC_PLACEHOLDER__"

    def apply_heading_style(self, paragraph, level: int):
        """
        Terapkan style Heading X dan pastikan outlineLvl terisi di awal pPr.
        """
        level = max(1, min(4, int(level)))
        style_name = f"Heading {level}"
        paragraph.style = style_name

        # Explicit outline level for TOC discovery (0 to 3 for Heading 1 to 4).
        # Keep element order stable by not forcing index 0 in pPr.
        p_pr = paragraph._p.get_or_add_pPr()
        outline_el = p_pr.find(qn("w:outlineLvl"))
        if outline_el is None:
            outline_el = OxmlElement("w:outlineLvl")
            p_pr.append(outline_el)
        outline_el.set(qn("w:val"), str(level - 1))

        enable_flow_control(paragraph, keep_next=True, keep_lines=True)
        return paragraph

    def insert_heading(self, text: str, level: int = 1):
        """
        Tambahkan satu paragraf heading dengan style native yang benar.
        """
        paragraph = self.document.add_paragraph()
        self.apply_heading_style(paragraph, level)
        
        # PENTING: Gunakan teks mentah, biarkan Word Style (Heading X)
        # yang mengatur font, ukuran, dan ketebalan secara native.
        clean_text = " ".join((text or "").replace("\n", " ").split()).strip()
        if clean_text:
            paragraph.add_run(clean_text)
        return paragraph

    def insert_caption(self, text: str):
        """
        Khusus untuk caption gambar agar masuk TOC (sebagai level 4).
        """
        return self.insert_heading(text, level=4)

    def _build_toc_field_paragraph(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.set(qn("w:dirty"), "true")
        instr_text.text = r'TOC \o "1-4" \h \z \u \t "Heading 1;1;Heading 2;2;Heading 3;3;Heading 4;4"'

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        fld_separate.set(qn("w:dirty"), "true")

        placeholder_run = OxmlElement("w:r")
        placeholder_text = OxmlElement("w:t")
        placeholder_text.text = "Daftar isi akan diperbarui saat dokumen dibuka."
        placeholder_run.append(placeholder_text)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        fld_end.set(qn("w:dirty"), "true")

        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_separate)
        run._r.append(placeholder_run)
        run._r.append(fld_end)
        return paragraph

    def add_toc(self, title: str = "DAFTAR ISI"):
        """
        Membuat daftar isi Word asli yang stabil dan kompatibel.
        """

        # ======================================================
        # JUDUL DAFTAR ISI
        # ======================================================
        title_para = self.document.add_paragraph()

        # PENTING: Judul "DAFTAR ISI" menggunakan style centere manual
        # agar tidak masuk ke dalam Daftar Isi itu sendiri.
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.upper())
        set_run_fonts(title_run, font_name=FONT_NAME, size=HEADING_SIZE, bold=True)

        title_para.paragraph_format.space_after = Pt(18)

        # ======================================================
        # FIELD TOC WORD ASLI
        # ======================================================
        return self._build_toc_field_paragraph()

    def insert_toc(self, title: str = "DAFTAR ISI"):
        return self.add_toc(title)

    def insert_toc_placeholder(self):
        marker = self.document.add_paragraph(self._placeholder_tag)
        marker.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return marker

    def finalize_toc(self, title: str = "DAFTAR ISI"):
        for paragraph in self.document.paragraphs:
            if (paragraph.text or "").strip() == self._placeholder_tag:
                marker_element = paragraph._p
                parent = marker_element.getparent()
                insert_at = parent.index(marker_element)

                self.add_toc(title=title)
                title_p = self.document.paragraphs[-2]._p
                toc_p = self.document.paragraphs[-1]._p

                parent.remove(title_p)
                parent.remove(toc_p)
                parent.insert(insert_at, title_p)
                parent.insert(insert_at + 1, toc_p)
                parent.remove(marker_element)
                return True
        return False