from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .docx_utils import set_cell_width, set_repeat_table_header, set_table_borders_none
from .paragraph_engine import enable_flow_control
from .styles import BODY_SIZE, TEXT_WIDTH_CM, apply_caption_style, format_run


class LayoutEngine:
    def __init__(self, document):
        self.document = document

    def add_invisible_table(self, rows: int, cols: int, widths_cm: list[float] | None = None):
        table = self.document.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders_none(table)
        if widths_cm:
            for row in table.rows:
                for index, cell in enumerate(row.cells):
                    set_cell_width(cell, widths_cm[index])
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        return table

    def fill_table_cell(self, cell, text: str = "", *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        enable_flow_control(paragraph, keep_lines=True)
        if text:
            format_run(paragraph.add_run(text), bold=bold, italic=False, size=BODY_SIZE)
        return paragraph

    def add_data_table(self, headers: list[str], rows: list[list], caption: str = ""):
        row_count = len(rows) + (1 if headers else 0)
        if row_count == 0:
            return None
        total_cols = max(len(headers), max((len(row) for row in rows), default=0))
        total_cols = max(total_cols, 1)
        widths = [TEXT_WIDTH_CM / total_cols for _ in range(total_cols)]
        table = self.document.add_table(rows=row_count, cols=total_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        current_row = 0
        if headers:
            header_row = table.rows[0]
            set_repeat_table_header(header_row)
            for col_index, value in enumerate(headers[:total_cols]):
                cell = header_row.cells[col_index]
                set_cell_width(cell, widths[col_index])
                self.fill_table_cell(cell, str(value), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            current_row = 1

        for row in rows:
            for col_index in range(total_cols):
                cell = table.rows[current_row].cells[col_index]
                set_cell_width(cell, widths[col_index])
                value = row[col_index] if col_index < len(row) else ""
                self.fill_table_cell(cell, str(value), align=WD_ALIGN_PARAGRAPH.LEFT)
            current_row += 1

        if caption.strip():
            paragraph = self.document.add_paragraph()
            apply_caption_style(paragraph)
            format_run(paragraph.add_run(caption.strip()), italic=False, bold=False)
        return table
