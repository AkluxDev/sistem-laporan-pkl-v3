"""
ImageEngine v3.0 — Penempatan Gambar Standar Akademik

Perbaikan:
- Nomor gambar otomatis: "Gambar 1. Keterangan"
- Spacing atas-bawah sesuai standar laporan PKL
- Gambar kecil di-scale UP ke lebar yang proporsional (bukan hanya shrink)
- Border ringan opsional untuk gambar isi laporan
- Gambar selalu center, caption selalu di bawah gambar (standar)
- Reset counter per-dokumen
"""

import io

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from .docx_utils import decode_image_stream, get_image_dimensions_cm


# Lebar teks halaman A4 dengan margin standar PKL (4cm kiri, 3cm kanan = 14cm)
TEXT_WIDTH_CM = 14.0
# Tinggi max gambar isi agar tidak melampaui 1 halaman
MAX_HEIGHT_CM = 17.0


class ImageEngine:
    """
    ImageEngine V3 — Standar Akademik untuk Laporan PKL.
    - Nomor caption otomatis (Gambar 1, Gambar 2, ...)
    - Scaling proporsional: gambar kecil di-upscale, gambar besar di-downscale
    - Spacing konsisten sesuai standar
    """

    def __init__(self, document):
        self.document = document
        self._figure_counter = 0  # Counter nomor gambar per dokumen

    def _next_figure_num(self) -> int:
        self._figure_counter += 1
        return self._figure_counter

    def insert_cover_image(
        self,
        image_source,
        max_width_cm: float = 6.0,
        max_height_cm: float = 6.0,
    ):
        """Sisipkan gambar cover (logo sekolah) — center, proporsional, tanpa caption."""
        if not image_source:
            return None

        stream = decode_image_stream(image_source)
        if not stream:
            return None

        w, h = self._fit_dimensions(stream, max_width_cm, max_height_cm, allow_upscale=False)

        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after  = Pt(12)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run()
        stream.seek(0)
        run.add_picture(stream, width=Cm(w), height=Cm(h))
        return paragraph

    def insert_content_image(
        self,
        image_source,
        caption: str = "",
        max_width_cm: float = TEXT_WIDTH_CM,
        max_height_cm: float = MAX_HEIGHT_CM,
        toc_entry_text: str = "",
    ):
        """
        Sisipkan gambar ke dalam isi laporan.

        Layout standar laporan PKL:
        - Spasi 12pt sebelum gambar
        - Gambar center, proporsional
        - Caption di bawah: "Gambar N. [keterangan]" (bold nomor, italic teks)
        - Spasi 12pt setelah caption
        """
        if not image_source:
            return None

        stream = decode_image_stream(image_source)
        if not stream:
            return None

        fig_num = self._next_figure_num()

        # Hitung dimensi yang pas
        w, h = self._fit_dimensions(stream, max_width_cm, max_height_cm, allow_upscale=True)

        # ── Paragraf gambar ─────────────────────────────────────────
        img_para = self.document.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.left_indent       = Cm(0)
        img_para.paragraph_format.first_line_indent = Cm(0)
        img_para.paragraph_format.space_before      = Pt(12)
        img_para.paragraph_format.space_after       = Pt(4)
        img_para.paragraph_format.line_spacing      = 1.0

        run = img_para.add_run()
        stream.seek(0)
        run.add_picture(stream, width=Cm(w), height=Cm(h))

        # ── Caption ─────────────────────────────────────────────────
        caption_text = caption.strip() if caption else ""
        full_caption = f"Gambar {fig_num}. {caption_text}" if caption_text else f"Gambar {fig_num}"

        cap_para = self.document.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.left_indent       = Cm(0)
        cap_para.paragraph_format.first_line_indent = Cm(0)
        cap_para.paragraph_format.space_before      = Pt(2)
        cap_para.paragraph_format.space_after       = Pt(12)
        cap_para.paragraph_format.line_spacing      = 1.0

        # Format caption: "Gambar N." bold, teks keterangan italic
        run_num = cap_para.add_run(f"Gambar {fig_num}.")
        run_num.bold   = True
        run_num.italic = False
        _set_run_font(run_num, size_pt=11)

        if caption_text:
            run_desc = cap_para.add_run(f" {caption_text}")
            run_desc.bold   = False
            run_desc.italic = True
            _set_run_font(run_desc, size_pt=11)

        # Daftarkan ke TOC engine jika tersedia
        if hasattr(self.document, '_toc_engine'):
            try:
                self.document._toc_engine.register_figure(fig_num, caption_text or "")
            except Exception:
                pass

        return img_para

    def insert_image_fit(
        self,
        image_source,
        max_width_cm: float,
        max_height_cm: float,
        paragraph=None,
    ):
        """Sisipkan gambar (misal TTD) ke paragraf tertentu agar pas di area tersedia."""
        if not image_source:
            return None

        stream = decode_image_stream(image_source)
        if not stream:
            return None

        w, h = self._fit_dimensions(stream, max_width_cm, max_height_cm, allow_upscale=False)

        if paragraph is None:
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        stream.seek(0)
        run.add_picture(stream, width=Cm(w), height=Cm(h))
        return paragraph

    # ── Internal helpers ──────────────────────────────────────────────

    def _fit_dimensions(
        self,
        stream: io.BytesIO,
        max_w: float,
        max_h: float,
        allow_upscale: bool = True,
    ) -> tuple[float, float]:
        """
        Hitung (width, height) dalam cm agar gambar proporsional di dalam batas max.

        allow_upscale=True  → gambar kecil di-scale naik hingga lebar halaman
        allow_upscale=False → gambar kecil dibiarkan ukuran aslinya (untuk cover/TTD)
        """
        dims = get_image_dimensions_cm(stream)

        if not dims:
            # Fallback jika PIL tidak tersedia: gunakan lebar 2/3 halaman
            fallback_w = min(max_w, TEXT_WIDTH_CM * 0.75)
            return fallback_w, fallback_w * 0.6  # rasio 5:3

        orig_w, orig_h = dims

        # Hitung rasio scale untuk masuk ke batas max
        scale = min(max_w / orig_w, max_h / orig_h)

        if not allow_upscale:
            # Jangan perbesar melebihi ukuran asli
            scale = min(scale, 1.0)
        else:
            # Upscale sampai mencapai lebar halaman, tapi tidak melebihi max
            # Minimal 60% lebar teks agar gambar tidak terlalu kecil
            min_scale = max(0.6 * TEXT_WIDTH_CM / orig_w, scale)
            # Ambil scale yang lebih besar (upscale), tapi tetap dalam batas max
            scale = min(min_scale, max_w / orig_w, max_h / orig_h)

        return orig_w * scale, orig_h * scale


def _set_run_font(run, size_pt: float = 11, font_name: str = "Times New Roman"):
    """Set font name dan size pada run."""
    run.font.name  = font_name
    run.font.size  = Pt(size_pt)
    # Set font untuk East Asian juga
    try:
        from docx.oxml.ns import qn as _qn
        rpr = run._r.get_or_add_rPr()
        rFonts = rpr.get_or_add_rFonts()
        rFonts.set(_qn('w:ascii'),    font_name)
        rFonts.set(_qn('w:hAnsi'),    font_name)
        rFonts.set(_qn('w:cs'),       font_name)
        rFonts.set(_qn('w:eastAsia'), font_name)
    except Exception:
        pass
