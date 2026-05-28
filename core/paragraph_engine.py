"""
paragraph_engine.py
===================
Engine untuk membuat heading dan paragraph dengan struktur XML yang benar.

PERBAIKAN KRITIS:
1. outlineLvl wajib ditulis ke XML untuk setiap heading
2. TIDAK ada run.font.bold manual pada heading (merusak style)
3. TIDAK ada style override setelah heading
4. Heading 1-4 semua menggunakan style native Word
5. Sub bab = Heading 2, Sub list = Heading 3, Caption = Heading 4
6. Tidak ada \\n pada teks heading

PENJELASAN MASALAH LAMA:
- python-docx terkadang tidak menulis outlineLvl ke XML
- Tanpa outlineLvl, Word tidak mengenali paragraph sebagai heading untuk TOC
- Run font manual (run.font.bold, run.font.name) MENIMPA style heading
- Word membaca style + outlineLvl untuk menentukan apakah masuk TOC
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .docx_utils import sanitize_paragraph_lines
from .styles import (
    BODY_SIZE,
    HEADING_SIZE,
    SUBHEADING_SIZE,
    apply_heading_style,
    apply_list_style,
    apply_normal_style,
    apply_signature_style,
    apply_subheading_style,
    format_run,
)


# ============================================================
# KONSTANTA OUTLINE LEVEL
# ============================================================

OUTLINE_LEVEL = {
    1: 0,   # Heading 1 = outlineLvl 0
    2: 1,   # Heading 2 = outlineLvl 1
    3: 2,   # Heading 3 = outlineLvl 2
    4: 3,   # Heading 4 = outlineLvl 3
}

HEADING_STYLE_MAP = {
    1: 'Heading 1',
    2: 'Heading 2',
    3: 'Heading 3',
    4: 'Heading 4',
}


# ============================================================
# FUNGSI INTI: Pastikan outlineLvl ada di pPr
# ============================================================

def _ensure_outline_level(paragraph, level):
    """
    Inject atau update outlineLvl di XML paragraph.
    
    MENGAPA PERLU:
    - python-docx tidak selalu menulis outlineLvl ke XML
    - Word WAJIB memiliki outlineLvl untuk memasukkan paragraph ke TOC
    - Fungsi ini memastikan outlineLvl selalu ada
    
    Args:
        paragraph: docx Paragraph object
        level: int 1-4 (heading level)
    """
    outline_value = OUTLINE_LEVEL.get(level, level - 1)
    
    pPr = paragraph._element.get_or_add_pPr()
    
    # Hapus outlineLvl lama jika ada
    old_outline = pPr.find(qn('w:outlineLvl'))
    if old_outline is not None:
        pPr.remove(old_outline)
    
    # Buat outlineLvl baru
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(outline_value))
    pPr.append(outlineLvl)


def _remove_conflicting_styles(paragraph):
    """
    Hapus elemen XML yang bisa bentrok dengan heading style.
    
    MENGAPA PERLU:
    - numPr (list numbering) pada heading merusak TOC detection
    - ind (indentation manual) bisa bentrok dengan heading style
    - Word membaca style heading dari pStyle, tapi list/indent bisa override
    """
    pPr = paragraph._element.get_or_add_pPr()
    
    # Hapus numPr jika ada (list numbering tidak boleh ada di heading)
    num_pr = pPr.find(qn('w:numPr'))
    if num_pr is not None:
        pPr.remove(num_pr)


def _clean_run_formatting(paragraph):
    """
    Bersihkan run-level formatting yang bisa merusak heading style.
    
    MENGAPA PERLU:
    - run.font.bold = True pada heading MENIMPA style heading
    - Ini menyebabkan heading terlihat benar secara visual
    - Tapi Word kadang gagal mendeteksi sebagai heading untuk TOC
    - Solusi: biarkan heading style yang menentukan font/bold/size
    
    CATATAN: Fungsi ini hanya menghapus rPr yang ada di dalam run heading.
    Font normal tidak terpengaruh.
    """
    for run in paragraph.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            # Hapus bold override jika ada
            bold_elem = rPr.find(qn('w:b'))
            if bold_elem is not None:
                rPr.remove(bold_elem)
            
            # Hapus font override jika ada
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                rPr.remove(rFonts)
            
            # Hapus size override jika ada
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                rPr.remove(sz)
            
            szCs = rPr.find(qn('w:szCs'))
            if szCs is not None:
                rPr.remove(szCs)


# ============================================================
# FUNGSI UTAMA: Buat Heading
# ============================================================

def add_heading(doc, text, level, clean_formatting=True):
    """
    Buat heading native Word dengan outlineLvl yang benar.
    
    ATURAN PENTING:
    - text TIDAK BOLEH mengandung \\n
    - level 1-4 saja
    - Jangan panggil run.font.bold setelah fungsi ini
    - Jangan ganti style setelah fungsi ini
    
    Args:
        doc: Document object
        text: str - teks heading (TANPA \\n)
        level: int - 1=BAB, 2=SubBab, 3=SubList, 4=Caption
        clean_formatting: bool - True untuk bersihkan run override
    
    Returns:
        paragraph object
    
    Contoh:
        add_heading(doc, "BAB I PENDAHULUAN", 1)
        add_heading(doc, "A. Latar Belakang", 2)
        add_heading(doc, "1. Instalasi Jaringan", 3)
        add_heading(doc, "Gambar 1. Topologi", 4)
    """
    # Validasi input
    if not text:
        raise ValueError("Teks heading tidak boleh kosong")
    
    if '\\n' in text or '\n' in text:
        # Bersihkan newline otomatis
        text = text.replace('\\n', ' ').replace('\n', ' ').strip()
    
    if level not in HEADING_STYLE_MAP:
        raise ValueError(f"Level heading harus 1-4, bukan {level}")
    
    # Buat paragraph dengan style heading yang tepat
    style_name = HEADING_STYLE_MAP[level]
    
    para = doc.add_paragraph(style=style_name)
    
    # Tambah teks sebagai run
    run = para.add_run(text)
    
    # Pastikan outlineLvl ada di XML
    _ensure_outline_level(para, level)
    
    # Hapus style yang bentrok
    _remove_conflicting_styles(para)
    
    # Bersihkan run formatting override (opsional tapi direkomendasikan)
    if clean_formatting:
        _clean_run_formatting(para)
    
    return para


# ============================================================
# FUNGSI: Buat Heading BAB (Heading 1)
# ============================================================

def add_bab(doc, nomor_bab, judul_bab):
    """
    Buat heading BAB sebagai Heading 1 native Word.
    
    Format: "BAB {nomor_bab} {judul_bab}"
    Contoh: add_bab(doc, "I", "PENDAHULUAN") → "BAB I PENDAHULUAN"
    
    WAJIB: Semua dalam satu baris, tanpa \\n
    """
    text = f"BAB {nomor_bab} {judul_bab}".strip()
    return add_heading(doc, text, level=1)


# ============================================================
# FUNGSI: Buat Sub Bab (Heading 2)
# ============================================================

def add_sub_bab(doc, label, judul):
    """
    Buat sub bab sebagai Heading 2 native Word.
    
    Format: "{label}. {judul}" atau "{label} {judul}"
    Contoh: add_sub_bab(doc, "A", "Latar Belakang") → "A. Latar Belakang"
    
    Sub bab WAJIB Heading 2 agar muncul di TOC.
    """
    if label:
        text = f"{label}. {judul}".strip()
    else:
        text = judul.strip()
    
    return add_heading(doc, text, level=2)


# ============================================================
# FUNGSI: Buat Sub List / Sub Sub Bab (Heading 3)
# ============================================================

def add_sub_list_heading(doc, nomor, judul):
    """
    Buat sub list sebagai Heading 3 native Word.
    
    Format: "{nomor}. {judul}"
    Contoh: add_sub_list_heading(doc, "1", "Instalasi Jaringan") → "1. Instalasi Jaringan"
    
    PENTING: Ini adalah HEADING TERPISAH, bukan paragraph list.
    Teks penjelasan dibuat sebagai paragraph biasa SETELAH heading ini.
    
    Sub list WAJIB Heading 3 agar muncul di TOC.
    """
    if nomor:
        text = f"{nomor}. {judul}".strip()
    else:
        text = judul.strip()
    
    return add_heading(doc, text, level=3)


# ============================================================
# FUNGSI: Buat Caption Gambar (Heading 4)
# ============================================================

def add_caption_heading(doc, nomor_gambar, keterangan):
    """
    Buat caption gambar sebagai Heading 4 native Word.
    
    Format: "Gambar {nomor}. {keterangan}"
    Contoh: add_caption_heading(doc, "1", "Topologi Jaringan")
    
    Caption WAJIB Heading 4 agar muncul di TOC.
    """
    text = f"Gambar {nomor_gambar}. {keterangan}".strip()
    return add_heading(doc, text, level=4)


# ============================================================
# FUNGSI: Buat Paragraph Biasa
# ============================================================

def add_paragraph(doc, text, style='Normal', alignment=None):
    """
    Buat paragraph teks biasa (bukan heading).
    
    Args:
        doc: Document object
        text: str - isi teks
        style: str - style paragraph (default 'Normal')
        alignment: WD_ALIGN_PARAGRAPH atau None
    
    Returns:
        paragraph object
    """
    para = doc.add_paragraph(text, style=style)
    
    if alignment is not None:
        para.alignment = alignment
    
    return para


# ============================================================
# FUNGSI: Buat List Item (BUKAN heading, paragraph biasa dengan numbering)
# ============================================================

def add_list_item(doc, text, level=0, list_style='List Paragraph'):
    """
    Buat list item sebagai paragraph biasa dengan numbering.
    
    PENTING: List item BUKAN heading.
    Jika konten perlu masuk TOC, gunakan add_sub_list_heading().
    
    Args:
        doc: Document object
        text: str - isi list
        level: int - indentation level (0=level pertama, 1=nested, dst)
        list_style: str - style untuk list
    
    Returns:
        paragraph object
    """
    try:
        para = doc.add_paragraph(text, style=list_style)
    except KeyError:
        para = doc.add_paragraph(text, style='Normal')
    
    # Set indentation berdasarkan level
    if level > 0:
        from docx.shared import Cm
        para.paragraph_format.left_indent = Cm(level * 1.0)
    
    return para


# ============================================================
# FUNGSI DEBUG: Cek semua heading di document
# ============================================================

def debug_headings(doc):
    """
    Debug helper: tampilkan semua heading dan outlineLvl mereka.
    
    Gunakan ini untuk verifikasi bahwa heading sudah benar.
    """
    print("\n=== DEBUG HEADINGS ===")
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "None"
        
        if 'Heading' in style_name:
            # Cek outlineLvl
            pPr = para._element.find(qn('w:pPr'))
            outline_val = "NOT FOUND"
            
            if pPr is not None:
                outline_elem = pPr.find(qn('w:outlineLvl'))
                if outline_elem is not None:
                    outline_val = outline_elem.get(qn('w:val'), 'NO VALUE')
            
            print(f"Para [{i}] Style='{style_name}' outlineLvl={outline_val} Text='{para.text[:50]}'")
    print("=== END DEBUG ===\n")

def enable_flow_control(paragraph, keep_next: bool = False, keep_lines: bool = False, widow_control: bool = True):
    pPr = paragraph._element.get_or_add_pPr()
    if keep_next:
        keepNext = pPr.find(qn('w:keepNext'))
        if keepNext is None:
            keepNext = OxmlElement('w:keepNext')
            pPr.append(keepNext)
    if keep_lines:
        keepLines = pPr.find(qn('w:keepLines'))
        if keepLines is None:
            keepLines = OxmlElement('w:keepLines')
            pPr.append(keepLines)
    if widow_control:
        widowControl = pPr.find(qn('w:widowControl'))
        if widowControl is None:
            widowControl = OxmlElement('w:widowControl')
            pPr.append(widowControl)


class ParagraphEngine:
    def __init__(self, document):
        self.document = document

    def add_body_text(self, text: str):
        paragraphs = []
        for block in sanitize_paragraph_lines(text):
            paragraph = self.document.add_paragraph()
            apply_normal_style(paragraph)
            enable_flow_control(paragraph, keep_lines=True)
            run = paragraph.add_run(block)
            format_run(run, size=BODY_SIZE, italic=False, bold=False)
            paragraphs.append(paragraph)
        return paragraphs

    def add_heading_bab(self, label: str, title: str | None = None):
        paragraph = self.document.add_paragraph()
        apply_heading_style(paragraph, level=1)
        enable_flow_control(paragraph, keep_next=True, keep_lines=True)
        
        heading_text = label.upper() if not title else f"{label.upper()} {title.upper()}"
        # PENTING: Jangan gunakan format_run agar Heading Style asli Word 
        # yang mengontrol font, ukuran, dan ketebalan.
        paragraph.add_run(heading_text.strip())
        return [paragraph]

    def add_section_heading(self, text: str):
        paragraph = self.document.add_paragraph()
        apply_heading_style(paragraph, level=1)
        enable_flow_control(paragraph, keep_next=True, keep_lines=True)
        
        # PENTING: Jangan gunakan format_run agar Heading Style asli Word yang mengontrol.
        paragraph.add_run(text.upper().strip())
        return paragraph

    def add_subheading(self, text: str, level: int = 2):
        paragraph = self.document.add_paragraph()
        apply_subheading_style(paragraph, level=level)
        enable_flow_control(paragraph, keep_next=True, keep_lines=True)
        
        # PENTING: Jangan gunakan format_run agar Heading Style asli Word yang mengontrol.
        paragraph.add_run(text.strip())
        return paragraph

    def add_centered_line(self, text: str, *, bold=False, size=None):
        paragraph = self.document.add_paragraph()
        apply_signature_style(paragraph)
        enable_flow_control(paragraph, keep_lines=True)
        format_run(paragraph.add_run(text), bold=bold, italic=False, size=size or BODY_SIZE)
        return paragraph

    def add_list_item(self, label: str, title: str, text: str, level: int = 0):
        paragraph = self.document.add_paragraph()
        apply_list_style(paragraph, level=level)
        enable_flow_control(paragraph, keep_lines=True)
        if label:
            format_run(paragraph.add_run(f"{label} "), bold=True, italic=False)
        if title:
            format_run(paragraph.add_run(f"{title} "), bold=True, italic=False)
        if text:
            format_run(paragraph.add_run(text), bold=False, italic=False)
        return paragraph