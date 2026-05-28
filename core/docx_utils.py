"""
docx_utils.py
=============
Utilities untuk manipulasi DOCX pada level XML.

PERBAIKAN KRITIS:
1. update_fields_setting() - menyimpan <w:updateFields w:val="true"/> ke settings.xml
2. ensure_all_fields_dirty() - pastikan semua field punya dirty=true
3. Fungsi ini WAJIB dipanggil sebelum doc.save()

PENJELASAN MENGAPA updateFields PERLU:
- python-docx tidak menyimpan updateFields ke settings.xml secara default
- Tanpa ini, Word tidak akan auto-update TOC saat file dibuka
- Harus diinjeksikan ke XML settings.xml di dalam DOCX secara manual

PENJELASAN MENGAPA dirty=true PERLU:
- Word menggunakan dirty flag untuk tahu field mana yang perlu di-recalculate
- Tanpa dirty=true, Word menganggap nilai field sudah benar
- Dengan dirty=true, Word tahu harus hitung ulang field tersebut
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import shutil
import os
import copy
import io
import base64
import datetime as _dt
import re
from typing import Any, Iterable, Optional
from docx.shared import Cm

try:
    from PIL import Image as PILImage
except ImportError:  # pragma: no cover
    PILImage = None


# ============================================================
# FUNGSI UTAMA: Aktifkan auto-update fields
# ============================================================

def enable_auto_update_fields(doc):
    """
    Aktifkan auto-update fields dengan menulis ke settings document.
    
    Ini adalah fungsi PALING PENTING untuk memastikan TOC dan nomor
    halaman ter-update otomatis saat dokumen dibuka.
    
    CARA KERJA:
    - Mengakses doc.settings._element (settings.xml)
    - Menambahkan <w:updateFields w:val="true"/>
    - Word membaca ini saat dokumen dibuka
    - Word langsung recalculate semua field yang dirty
    
    WAJIB dipanggil sebelum doc.save()
    
    Args:
        doc: Document object
    
    Returns:
        doc (untuk chaining)
    
    Contoh:
        doc = Document()
        # ... tambah konten ...
        enable_auto_update_fields(doc)
        doc.save("laporan.docx")
    """
    settings = doc.settings.element
    
    # Namespace yang digunakan
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # Cek apakah updateFields sudah ada
    update_fields_tag = f'{{{w_ns}}}updateFields'
    existing = settings.find(update_fields_tag)
    
    if existing is not None:
        # Update nilai yang sudah ada
        existing.set(f'{{{w_ns}}}val', 'true')
    else:
        # Buat elemen baru
        updateFields = OxmlElement('w:updateFields')
        updateFields.set(qn('w:val'), 'true')
        settings.insert(0, updateFields)
    
    return doc


# ============================================================
# FUNGSI: Pastikan semua field di document punya dirty=true
# ============================================================

def ensure_all_fields_dirty(doc):
    """
    Scan semua field di document dan pastikan semua dirty=true.
    
    Field yang di-scan:
    - fldChar (begin, separate, end)
    - Field di dalam paragraphs (body)
    - Field di dalam header/footer
    
    MENGAPA PERLU:
    - Jika ada field yang tidak dirty, Word tidak akan update field tersebut
    - Ini menyebabkan nomor halaman atau TOC tidak ter-update
    - Fungsi ini adalah safety net untuk memastikan semua field dirty
    
    Args:
        doc: Document object
    
    Returns:
        int: jumlah field yang di-update ke dirty=true
    """
    count = 0
    
    # Scan semua paragraph di body
    for para in doc.paragraphs:
        count += _make_paragraph_fields_dirty(para)
    
    # Scan header dan footer di semua section
    for section in doc.sections:
        # Header
        if not section.header.is_linked_to_previous:
            for para in section.header.paragraphs:
                count += _make_paragraph_fields_dirty(para)
        
        # Footer
        if not section.footer.is_linked_to_previous:
            for para in section.footer.paragraphs:
                count += _make_paragraph_fields_dirty(para)
    
    # Scan table cells jika ada
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += _make_paragraph_fields_dirty(para)
    
    return count


def _make_paragraph_fields_dirty(paragraph):
    """
    Pastikan semua fldChar dalam paragraph punya dirty=true.
    
    Returns:
        int: jumlah fldChar yang di-update
    """
    count = 0
    
    for child in paragraph._element.iter():
        if child.tag == qn('w:fldChar'):
            # Set dirty=true
            child.set(qn('w:dirty'), 'true')
            count += 1
    
    return count


# ============================================================
# FUNGSI: Tambah compatibility settings untuk Word
# ============================================================

def add_compatibility_settings(doc):
    """
    Tambah compatibility settings untuk meningkatkan stabilitas auto-update.
    
    Settings yang ditambahkan:
    - compatibilityMode: word2013 (OOXML modern)
    - Ini membantu Word menginterpretasikan field dengan benar
    
    Args:
        doc: Document object
    
    Returns:
        doc
    """
    settings = doc.settings.element
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    # Cek apakah compat sudah ada
    compat_tag = f'{{{w_ns}}}compat'
    compat = settings.find(compat_tag)
    
    if compat is None:
        compat = OxmlElement('w:compat')
        settings.append(compat)
    
    # Tambah compatSetting untuk Word 2013
    compat_setting = OxmlElement('w:compatSetting')
    compat_setting.set(qn('w:name'), 'compatibilityMode')
    compat_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
    compat_setting.set(qn('w:val'), '15')  # Word 2013
    
    # Cek apakah sudah ada
    existing_setting = None
    for child in compat:
        if (child.tag == qn('w:compatSetting') and 
            child.get(qn('w:name')) == 'compatibilityMode'):
            existing_setting = child
            break
    
    if existing_setting is not None:
        existing_setting.set(qn('w:val'), '15')
    else:
        compat.append(compat_setting)
    
    return doc


# ============================================================
# FUNGSI: Apply semua fix sekaligus (shortcut)
# ============================================================

def apply_all_docx_fixes(doc):
    """
    Shortcut: apply semua perbaikan sekaligus.
    
    Panggil fungsi ini SEBELUM doc.save():
    
        apply_all_docx_fixes(doc)
        doc.save("laporan.docx")
    
    Yang dilakukan:
    1. Enable auto-update fields di settings
    2. Pastikan semua field punya dirty=true
    3. Tambah compatibility settings
    
    Args:
        doc: Document object
    
    Returns:
        doc
    """
    # 1. Enable updateFields di settings.xml
    enable_auto_update_fields(doc)
    
    # 2. Pastikan semua field dirty
    dirty_count = ensure_all_fields_dirty(doc)
    
    # 3. Tambah compatibility settings
    add_compatibility_settings(doc)
    
    return doc


# ============================================================
# FUNGSI: Post-process DOCX file untuk inject updateFields
# ============================================================

def post_process_docx(input_path, output_path=None):
    """
    Post-process file DOCX yang sudah disimpan untuk memastikan
    updateFields ada di settings.xml.
    
    Ini adalah solusi alternatif jika enable_auto_update_fields() 
    tidak bekerja karena alasan tertentu.
    
    Cara kerja:
    1. Buka DOCX sebagai ZIP
    2. Extract settings.xml
    3. Inject <w:updateFields w:val="true"/>
    4. Simpan kembali ke DOCX
    
    Args:
        input_path: str - path ke file DOCX input
        output_path: str - path output (None = overwrite input)
    
    Returns:
        str: path file output
    """
    if output_path is None:
        output_path = input_path
    
    # Baca DOCX sebagai ZIP
    tmp_path = input_path + '.tmp'
    
    with zipfile.ZipFile(input_path, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                
                if item.filename == 'word/settings.xml':
                    # Inject updateFields ke settings.xml
                    data = _inject_update_fields(data)
                
                elif item.filename == 'word/document.xml':
                    # Pastikan semua fldChar punya dirty=true
                    data = _inject_dirty_fields(data)
                
                zout.writestr(item, data)
    
    # Ganti file asli
    if os.path.exists(output_path):
        os.remove(output_path)
    os.rename(tmp_path, output_path)
    
    return output_path


def _inject_update_fields(settings_xml_bytes):
    """
    Inject <w:updateFields w:val="true"/> ke settings.xml bytes.
    
    Returns:
        bytes: settings.xml yang sudah dimodifikasi
    """
    try:
        # Parse XML
        tree = etree.fromstring(settings_xml_bytes)
        
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        
        # Cari updateFields
        update_fields_tag = f'{{{w_ns}}}updateFields'
        existing = tree.find(update_fields_tag)
        
        if existing is not None:
            existing.set(f'{{{w_ns}}}val', '1')
        else:
            # Buat elemen baru
            update_fields = etree.SubElement(tree, update_fields_tag)
            update_fields.set(f'{{{w_ns}}}val', '1')
        
        # Serialize kembali
        return etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
    
    except Exception as e:
        # Jika gagal parse, kembalikan data asli
        return settings_xml_bytes


def _inject_dirty_fields(document_xml_bytes):
    """
    Inject dirty=true ke semua fldChar di document.xml bytes.
    
    Returns:
        bytes: document.xml yang sudah dimodifikasi
    """
    try:
        # Parse XML
        tree = etree.fromstring(document_xml_bytes)
        
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        fldChar_tag = f'{{{w_ns}}}fldChar'
        dirty_attr = f'{{{w_ns}}}dirty'
        
        # Scan semua fldChar
        for elem in tree.iter(fldChar_tag):
            elem.set(dirty_attr, 'true')
        
        # Serialize kembali
        return etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)
    
    except Exception as e:
        # Jika gagal parse, kembalikan data asli
        return document_xml_bytes


# ============================================================
# FUNGSI: Verifikasi settings.xml sudah benar
# ============================================================

def verify_docx_settings(docx_path):
    """
    Debug helper: verifikasi settings.xml di dalam DOCX.
    
    Cek:
    1. updateFields ada dan val=true
    2. Berapa banyak fldChar yang dirty
    
    Args:
        docx_path: str - path ke file DOCX
    
    Returns:
        dict berisi hasil verifikasi
    """
    result = {
        'updateFields_found': False,
        'updateFields_val': None,
        'fldChar_total': 0,
        'fldChar_dirty': 0,
        'headings_found': [],
    }
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            # Cek settings.xml
            if 'word/settings.xml' in z.namelist():
                settings_data = z.read('word/settings.xml')
                tree = etree.fromstring(settings_data)
                
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                update_tag = f'{{{w_ns}}}updateFields'
                val_attr = f'{{{w_ns}}}val'
                
                elem = tree.find(update_tag)
                if elem is not None:
                    result['updateFields_found'] = True
                    result['updateFields_val'] = elem.get(val_attr)
            
            # Cek document.xml
            if 'word/document.xml' in z.namelist():
                doc_data = z.read('word/document.xml')
                tree = etree.fromstring(doc_data)
                
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                fldChar_tag = f'{{{w_ns}}}fldChar'
                dirty_attr = f'{{{w_ns}}}dirty'
                pStyle_tag = f'{{{w_ns}}}pStyle'
                val_attr = f'{{{w_ns}}}val'
                
                for elem in tree.iter(fldChar_tag):
                    result['fldChar_total'] += 1
                    if elem.get(dirty_attr) == 'true':
                        result['fldChar_dirty'] += 1
                
                # Cek headings
                for elem in tree.iter(pStyle_tag):
                    val = elem.get(val_attr, '')
                    if 'Heading' in val or 'heading' in val.lower():
                        result['headings_found'].append(val)
    
    except Exception as e:
        result['error'] = str(e)
    
    return result


# ============================================================
# FUNGSI LAMA YANG DIKEMBALIKAN (dari sistem.txt)
# ============================================================

BULAN_ID = [
    "",
    "Januari",
    "Februari",
    "Maret",
    "April",
    "Mei",
    "Juni",
    "Juli",
    "Agustus",
    "September",
    "Oktober",
    "November",
    "Desember",
]

def cm_to_twips(value_cm: float) -> int:
    return int(value_cm * 567)


def set_run_fonts(run, font_name: str, size=None, bold=False, italic=False, underline=False):
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    run.underline = underline
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def add_field_run(paragraph, instruction: str, font_name: str, size=None, display_text: str = ""):
    run = paragraph.add_run()
    set_run_fonts(run, font_name=font_name, size=size)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.set(qn("w:dirty"), "true")
    instr.text = f" {instruction} "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_separate.set(qn("w:dirty"), "true")
    text = OxmlElement("w:t")
    text.text = display_text
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end.set(qn("w:dirty"), "true")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    if display_text:
        run._r.append(text)
    run._r.append(fld_end)
    return run


def add_tc_entry(paragraph, entry_text: str, level: int = 3, identifier: str = "C", font_name: str = "Times New Roman", size=None):
    safe_text = (entry_text or "").replace('"', "'").strip()
    if not safe_text:
        return None
    instruction = f'TC "{safe_text}" \\l {level} \\f {identifier}'
    return add_field_run(paragraph, instruction, font_name=font_name, size=size)


def sanitize_inline_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (list, tuple, set)):
        parts = [sanitize_inline_text(item) for item in value]
        return " ".join(part for part in parts if part).strip()
    text = str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def sanitize_paragraph_lines(value: Any) -> list[str]:
    if value is None:
        return []
    
    raw_text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    
    lines = raw_text.split('\n')
    grouped_blocks = []
    current_block = []
    
    for line in lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            if current_block:
                grouped_blocks.append(" ".join(current_block))
                current_block = []
            continue
            
        if current_block:
            prev_line = current_block[-1].strip()
            
            starts_with_list = bool(re.search(r'^([0-9]+\.|[a-zA-Z]\.|-|\u2022)\s', line_stripped))
            ends_with_punc = bool(re.search(r'[.!?:"\']$', prev_line))
            starts_with_cap = bool(re.search(r'^[A-Z0-9"\'\u201C\u201D]', line_stripped))
            
            if starts_with_list or (ends_with_punc and starts_with_cap):
                grouped_blocks.append(" ".join(current_block))
                current_block = [line_stripped]
            else:
                current_block.append(line_stripped)
        else:
            current_block.append(line_stripped)
            
    if current_block:
        grouped_blocks.append(" ".join(current_block))
        
    final_blocks = [re.sub(r"[ \t]+", " ", block).strip() for block in grouped_blocks if block]
    return final_blocks


def ensure_update_fields_on_open(document):
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def coerce_dict_list(items: Any) -> list[dict]:
    if not isinstance(items, list):
        return []
    return [item for item in items if isinstance(item, dict)]


def take_first_non_empty(values: Iterable[Any]) -> str:
    for value in values:
        text = sanitize_inline_text(value)
        if text:
            return text
    return ""


def _normalize_image_bytes(img_bytes: bytes) -> Optional[io.BytesIO]:
    """Konversi bytes gambar ke BytesIO JPEG/PNG yang kompatibel dengan docx & reportlab."""
    if not img_bytes:
        return None
    try:
        if PILImage:
            from io import BytesIO
            img = PILImage.open(BytesIO(img_bytes))
            # Konversi format yang tidak didukung Word/ReportLab ke PNG
            if img.format not in ('JPEG', 'PNG', 'GIF', 'BMP'):
                if img.mode in ('RGBA', 'P', 'LA'):
                    bg = PILImage.new('RGBA', img.size, (255, 255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    if img.mode in ('RGBA', 'LA'):
                        bg.paste(img, mask=img.split()[-1])
                    else:
                        bg.paste(img)
                    img = bg.convert('RGB')
                elif img.mode not in ('RGB', 'L'):
                    img = img.convert('RGB')
                out = BytesIO()
                img.save(out, format='PNG')
                return io.BytesIO(out.getvalue())
            # RGBA/P PNG → komposit di atas putih agar Word tidak error
            if img.mode in ('RGBA', 'P', 'LA'):
                bg = PILImage.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                alpha = img.split()[-1] if img.mode in ('RGBA', 'LA') else None
                if alpha:
                    bg.paste(img.convert('RGB'), mask=alpha)
                else:
                    bg.paste(img)
                out = BytesIO()
                bg.save(out, format='PNG')
                return io.BytesIO(out.getvalue())
        return io.BytesIO(img_bytes)
    except Exception:
        return None


def decode_image_stream(image_value: Any) -> Optional[io.BytesIO]:
    """
    Decode berbagai format sumber gambar ke BytesIO yang siap pakai:
    - data URI (data:image/...;base64,...)
    - pure base64 string
    - file path (str)
    - bytes / BytesIO langsung
    """
    if image_value is None:
        return None

    # Sudah BytesIO — reset position
    if isinstance(image_value, io.BytesIO):
        image_value.seek(0)
        return image_value

    # Raw bytes
    if isinstance(image_value, (bytes, bytearray)):
        return _normalize_image_bytes(bytes(image_value))

    if not isinstance(image_value, str):
        return None

    image_value = image_value.strip()
    if not image_value:
        return None

    # data URI: data:image/png;base64,...
    if image_value.startswith("data:"):
        try:
            _, payload = image_value.split(",", 1)
            img_bytes = base64.b64decode(payload + "==")  # padding aman
            return _normalize_image_bytes(img_bytes)
        except Exception:
            return None

    # File path di disk
    if os.path.isfile(image_value):
        try:
            with open(image_value, "rb") as f:
                img_bytes = f.read()
            return _normalize_image_bytes(img_bytes)
        except Exception:
            return None

    # Coba decode sebagai pure base64 (tanpa data: prefix)
    try:
        img_bytes = base64.b64decode(image_value + "==")
        # Validasi: harus dimulai dengan magic bytes gambar
        if img_bytes[:4] in (b'\xff\xd8\xff\xe0', b'\xff\xd8\xff\xe1',  # JPEG
                              b'\x89PNG',  # PNG
                              b'GIF8',     # GIF
                              b'BM'):      # BMP
            return _normalize_image_bytes(img_bytes)
    except Exception:
        pass

    return None

def get_image_dimensions_cm(image_source: io.BytesIO) -> Optional[tuple[float, float]]:
    """
    Hitung dimensi gambar dalam cm berdasarkan piksel aktual.
    Selalu gunakan 96 DPI standar agar hasilnya konsisten di Word dan PDF.
    """
    if PILImage is None or not isinstance(image_source, io.BytesIO):
        return None
    try:
        image_source.seek(0)
        with PILImage.open(image_source) as image:
            STANDARD_DPI = 96.0
            width_cm  = (image.width  / STANDARD_DPI) * 2.54
            height_cm = (image.height / STANDARD_DPI) * 2.54
        image_source.seek(0)
        return width_cm, height_cm
    except Exception:
        try:
            image_source.seek(0)
        except Exception:
            pass
        return None

def format_indonesian_date(date_value: str) -> str:
    date_value = sanitize_inline_text(date_value)
    if not date_value:
        return ""
    try:
        dt = _dt.datetime.strptime(date_value, "%Y-%m-%d")
        return f"{dt.day} {BULAN_ID[dt.month]} {dt.year}"
    except ValueError:
        return date_value

def clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        element.remove(child)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, edge_data in kwargs.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_borders_none(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                left={"val": "nil"},
                bottom={"val": "nil"},
                right={"val": "nil"},
            )


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = tr_pr.find(qn("w:tblHeader"))
    if repeat is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(cm_to_twips(width_cm)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Cm(width_cm)