"""
page_engine.py
==============
Engine untuk section breaks dan page numbering Word.

PERBAIKAN KRITIS:
1. Cover = section 1, tanpa nomor halaman
2. Halaman awal (pengesahan, kata pengantar, daftar isi) = section 2, romawi (i, ii, iii)
3. Isi laporan = section 3, angka (1, 2, 3)
4. Setiap section harus disconnect dari previous (linkToPrevious = false)
5. Field PAGE dan NUMPAGES WAJIB dirty=true

PENJELASAN TEKNIS:
- Section break di Word dikontrol melalui sectPr di dalam paragraph terakhir setiap section
- sectPr dari section sebelumnya diletakkan di pPr paragraph terakhir section tersebut
- sectPr terakhir dokumen ada di body langsung
- Header/footer terhubung via relationship
- linkToPrevious=false WAJIB agar nomor halaman tidak bocor antar section
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from .docx_utils import add_field_run, clear_paragraph
from .styles import FONT_NAME, MARGIN_BOTTOM_CM, MARGIN_LEFT_CM, MARGIN_RIGHT_CM, MARGIN_TOP_CM
import copy


# ============================================================
# KONSTANTA
# ============================================================

# Format nomor halaman
PAGE_FORMAT_ROMAN = 'lowerRoman'   # i, ii, iii
PAGE_FORMAT_ARABIC = 'decimal'      # 1, 2, 3
PAGE_FORMAT_NONE = None             # Tanpa nomor halaman


# ============================================================
# FUNGSI HELPER: Buat field PAGE atau NUMPAGES
# ============================================================

def _make_page_field(field_type='PAGE'):
    """
    Buat run berisi field PAGE atau NUMPAGES dengan dirty=true.
    
    WAJIB dirty=true agar Word auto-update nomor halaman.
    
    Args:
        field_type: 'PAGE' atau 'NUMPAGES'
    
    Returns:
        list of XML elements untuk run
    """
    elements = []
    
    # Run 1: fldChar begin
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    fldChar1.set(qn('w:dirty'), 'true')
    r1.append(fldChar1)
    elements.append(r1)
    
    # Run 2: instrText
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = f' {field_type} '
    r2.append(instrText)
    elements.append(r2)
    
    # Run 3: fldChar separate
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar2.set(qn('w:dirty'), 'true')
    r3.append(fldChar2)
    elements.append(r3)
    
    # Run 4: placeholder value
    r4 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '1'
    r4.append(t)
    elements.append(r4)
    
    # Run 5: fldChar end
    r5 = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    fldChar3.set(qn('w:dirty'), 'true')
    r5.append(fldChar3)
    elements.append(r5)
    
    return elements


# ============================================================
# FUNGSI HELPER: Buat sectPr
# ============================================================

def _build_sectPr(page_format=None, start_value=1, link_to_previous=False):
    """
    Buat sectPr XML element untuk section properties.
    
    Args:
        page_format: 'lowerRoman', 'decimal', atau None (tanpa nomor)
        start_value: int - nomor halaman pertama
        link_to_previous: bool - True jika footer/header link ke section sebelumnya
    
    Returns:
        sectPr XML element
    """
    sectPr = OxmlElement('w:sectPr')
    
    # Page size A4
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')   # A4 width dalam twips
    pgSz.set(qn('w:h'), '16838')   # A4 height dalam twips
    sectPr.append(pgSz)
    
    # Margin (4cm top, 3cm bottom, 4cm left, 3cm right - standar nasional)
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '2268')     # 4cm
    pgMar.set(qn('w:right'), '1701')   # 3cm
    pgMar.set(qn('w:bottom'), '1701')  # 3cm
    pgMar.set(qn('w:left'), '2268')    # 4cm
    pgMar.set(qn('w:header'), '850')  # 1.5cm
    pgMar.set(qn('w:footer'), '850')   # 1.5cm
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)
    
    # Page numbering
    if page_format is not None:
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), page_format)
        pgNumType.set(qn('w:start'), str(start_value))
        sectPr.append(pgNumType)
    
    return sectPr


# ============================================================
# FUNGSI: Setup Footer dengan nomor halaman
# ============================================================

def _add_footer_to_sectPr(doc, sectPr, page_format, relationship_id=None):
    """
    Tambah footer reference ke sectPr.
    
    CATATAN: Footer harus sudah ada di document sebelum dipanggil.
    Gunakan doc.sections[-1].footer untuk mengakses footer.
    """
    # Footer reference
    footerRef = OxmlElement('w:footerReference')
    footerRef.set(qn('w:type'), 'default')
    
    if relationship_id:
        footerRef.set(qn('r:id'), relationship_id)
    
    sectPr.append(footerRef)
    
    return sectPr


# ============================================================
# FUNGSI: Tambah page number di footer menggunakan docx API
# ============================================================

def add_page_number_to_footer(section, page_format='decimal', center=True):
    """
    Tambah nomor halaman di footer section.
    
    Menggunakan python-docx API untuk footer, kemudian inject field PAGE.
    
    Args:
        section: docx Section object (dari doc.sections[-1])
        page_format: 'decimal' atau 'lowerRoman' (hanya untuk display,
                     format sebenarnya dikontrol di sectPr pgNumType)
        center: bool - True untuk center alignment
    
    Returns:
        footer paragraph
    """
    footer = section.footer
    
    # Hapus paragraphs yang ada
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    # Buat paragraph baru untuk footer
    footer_para = footer.add_paragraph()
    
    if center:
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Inject field PAGE dengan dirty=true
    p_elem = footer_para._element
    
    field_runs = _make_page_field('PAGE')
    for run_elem in field_runs:
        p_elem.append(run_elem)
    
    return footer_para


# ============================================================
# FUNGSI: Tambah section break dengan section properties
# ============================================================

def add_section_break(doc, page_format=None, start_value=1, break_type='nextPage'):
    """
    Tambah section break setelah paragraph terakhir dokumen saat ini.
    
    CARA KERJA:
    - Paragraph terakhir mendapat sectPr di dalam pPr-nya
    - sectPr ini mendefinisikan properties section SEBELUMNYA
    - Bukan section yang akan datang
    
    Args:
        doc: Document object
        page_format: 'lowerRoman', 'decimal', atau None
        start_value: int - mulai dari nomor berapa
        break_type: 'nextPage', 'evenPage', 'oddPage', 'continuous'
    
    Returns:
        sectPr element yang ditambahkan
    """
    # Ambil paragraph terakhir
    last_para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    
    pPr = last_para._element.get_or_add_pPr()
    
    # Hapus sectPr lama jika ada
    old_sectPr = pPr.find(qn('w:sectPr'))
    if old_sectPr is not None:
        pPr.remove(old_sectPr)
    
    # Buat sectPr baru
    sectPr = _build_sectPr(page_format=page_format, start_value=start_value)
    
    # Tambah type break
    if break_type != 'continuous':
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), break_type)
        sectPr.insert(0, sectType)
    
    pPr.append(sectPr)
    
    return sectPr


# ============================================================
# FUNGSI UTAMA: Setup complete page numbering system
# ============================================================

def setup_page_numbering(doc):
    """
    Setup sistem penomoran halaman lengkap.
    
    STRUKTUR YANG DIBUAT:
    
    Section 1 (Cover):
    - Tanpa nomor halaman
    - Hanya halaman cover
    
    Section 2 (Halaman Awal):
    - Format: i, ii, iii
    - Mulai dari i
    - Termasuk: Pengesahan, Kata Pengantar, Daftar Isi
    
    Section 3 (Isi Laporan):
    - Format: 1, 2, 3
    - Mulai dari 1
    - Termasuk: BAB I, dst
    
    CARA PENGGUNAAN:
    
    Panggil di doc_generator.py setelah semua konten ditambahkan:
    
        setup_page_numbering(doc)
    
    CATATAN PENTING:
    - Fungsi ini mengasumsikan konten sudah ada
    - Section break harus ditambahkan secara manual di tempat yang tepat
    - Gunakan add_section_break_marker() untuk menandai pergantian section
    
    Returns:
        doc dengan page numbering yang sudah disetup
    """
    sections = doc.sections
    
    if len(sections) >= 1:
        # Section 1: Cover - tanpa nomor halaman
        s1 = sections[0]
        _configure_section_no_page_number(s1)
    
    if len(sections) >= 2:
        # Section 2: Halaman awal - romawi
        s2 = sections[1]
        _configure_section_roman(s2)
    
    if len(sections) >= 3:
        # Section 3 dst: Isi laporan - angka
        for i in range(2, len(sections)):
            _configure_section_arabic(sections[i])
    
    return doc


def _configure_section_no_page_number(section):
    """
    Setup section tanpa nomor halaman (untuk cover).
    """
    # Hapus header dan footer
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    
    # Bersihkan footer
    footer = section.footer
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)
    
    # Tambah paragraph kosong di footer
    footer.add_paragraph()


def _configure_section_roman(section):
    """
    Setup section dengan nomor halaman romawi kecil (i, ii, iii).
    """
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    
    # Setup sectPr untuk roman numeral
    sectPr = section._sectPr
    
    # Update atau tambah pgNumType
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    
    pgNumType.set(qn('w:fmt'), 'lowerRoman')
    pgNumType.set(qn('w:start'), '1')
    
    # Tambah nomor halaman di footer
    add_page_number_to_footer(section, 'lowerRoman')


def _configure_section_arabic(section):
    """
    Setup section dengan nomor halaman angka (1, 2, 3).
    """
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    
    # Setup sectPr untuk decimal
    sectPr = section._sectPr
    
    # Update atau tambah pgNumType
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    
    # Tambah nomor halaman di footer
    add_page_number_to_footer(section, 'decimal')


# ============================================================
# FUNGSI: Tambah section break marker (paragraph kosong)
# ============================================================

def add_section_break_marker(doc, page_format, start_value=1):
    """
    Tambah paragraph khusus sebagai section break marker.
    
    Ini adalah cara yang benar untuk membuat section baru di python-docx.
    
    Cara penggunaan:
    
        # Setelah cover, sebelum konten romawi:
        add_section_break_marker(doc, 'lowerRoman', 1)
        
        # Setelah daftar isi, sebelum BAB I:
        add_section_break_marker(doc, 'decimal', 1)
    
    Args:
        doc: Document object
        page_format: 'lowerRoman' atau 'decimal' atau None
        start_value: int - mulai dari berapa
    
    Returns:
        paragraph yang dijadikan section break marker
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Buat paragraph kosong
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    
    # Tambah sectPr ke dalam pPr
    pPr = para._element.get_or_add_pPr()
    
    sectPr = OxmlElement('w:sectPr')
    
    # Section type: next page
    sectType = OxmlElement('w:type')
    sectType.set(qn('w:val'), 'nextPage')
    sectPr.append(sectType)
    
    # Page size A4
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    sectPr.append(pgSz)
    
    # Margin
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '2268')
    pgMar.set(qn('w:right'), '1701')
    pgMar.set(qn('w:bottom'), '1701')
    pgMar.set(qn('w:left'), '2268')
    pgMar.set(qn('w:header'), '850')  # 1.5cm
    pgMar.set(qn('w:footer'), '850')   # 1.5cm
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)
    
    # Page numbering format
    if page_format is not None:
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), page_format)
        pgNumType.set(qn('w:start'), str(start_value))
        sectPr.append(pgNumType)
    
    pPr.append(sectPr)
    
    return para


class PageEngine:
    def __init__(self, document):
        self.document = document
        self._configure_section(self.document.sections[0])

    def configure_base_section(self):
        self._configure_section(self.document.sections[0])
        return self.document.sections[0]

    def add_section(self, start=WD_SECTION_START.NEW_PAGE):
        section = self.document.add_section(start)
        self._configure_section(section)
        return section

    def add_page_break(self):
        self.document.add_page_break()

    def _configure_section(self, section):
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(MARGIN_LEFT_CM)
        section.top_margin = Cm(MARGIN_TOP_CM)
        section.right_margin = Cm(MARGIN_RIGHT_CM)
        section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        section.different_first_page_header_footer = False
        return section

    def setup_cover_section(self):
        section = self.document.sections[0]
        section.different_first_page_header_footer = True
        self._clear_footer(section.footer)
        self._clear_footer(section.first_page_footer)
        self._set_page_number_type(section, start=1, fmt="decimal")
        return section

    def setup_preliminary_section(self, section, start=1):
        """Setup section untuk Kata Pengantar & Daftar Isi (i, ii, iii...)."""
        section.different_first_page_header_footer = False
        self._attach_footer(section, page_format="lowerRoman", start=start)

    def setup_body_section(self, section, start=1):
        """Setup section untuk konten utama BAB I dst (1, 2, 3...)."""
        section.different_first_page_header_footer = True
        self._attach_footer(section, page_format="decimal", start=start, center=True, is_first_page=True)
        self._attach_header(section, page_format="decimal", center=False)
        self._clear_footer(section.footer)

    def _attach_footer(self, section, page_format: str, start: int, center: bool = True, is_first_page: bool = False):
        """Sisipkan nomor halaman menggunakan field PAGE yang dinamis."""
        footer = section.first_page_footer if is_first_page else section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        
        # Injeksi field PAGE secara manual melalui XML agar stabil
        run = paragraph.add_run()
        
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")
        
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.set(qn("w:dirty"), "true")  # Tandai instruksi sebagai dirty
        instr.text = " PAGE "

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        fld_separate.set(qn("w:dirty"), "true")

        text = OxmlElement("w:t")
        text.text = "1"

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        fld_end.set(qn("w:dirty"), "true")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_separate)
        run._r.append(text)
        run._r.append(fld_end)
        
        # Set tipe penomoran dan angka mulai pada property section
        self._set_page_number_type(section, start=start, fmt=page_format)

    def _attach_header(self, section, page_format: str, center: bool = False):
        """Sisipkan nomor halaman di header untuk halaman berikutnya."""
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        
        run = paragraph.add_run()
        
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")
        
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.set(qn("w:dirty"), "true")
        instr.text = " PAGE "

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        fld_separate.set(qn("w:dirty"), "true")

        text = OxmlElement("w:t")
        text.text = "1"

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        fld_end.set(qn("w:dirty"), "true")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_separate)
        run._r.append(text)
        run._r.append(fld_end)

    def _clear_footer(self, footer):
        footer.is_linked_to_previous = False
        for paragraph in footer.paragraphs:
            clear_paragraph(paragraph)

    def _set_page_number_type(self, section, start: int, fmt: str):
        sect_pr = section._sectPr
        page_num = sect_pr.find(qn("w:pgNumType"))
        if page_num is None:
            page_num = OxmlElement("w:pgNumType")
            sect_pr.append(page_num)
        page_num.set(qn("w:start"), str(start))
        page_num.set(qn("w:fmt"), fmt)
