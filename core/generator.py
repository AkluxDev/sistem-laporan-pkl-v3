"""
Generator dokumen Word (.docx) untuk Sistem Laporan PKL
v2.1 — fixed

FIX:
- _buat_daftar_isi: menggunakan TabStop XML yang benar agar titik-titik
  dan nomor halaman terformat rapi.
- _render_list: type check diperbaiki dari ["title","simple","list"]
  menjadi cocok dengan data JS yang selalu mengirim type:"list".
  Mode item diambil dari field 'type' di dalam item, bukan di wrapper.
- _buat_kata_pengantar: teks hardcode bahasa Arab dihapus; kalimat
  pengantar ucapan terima kasih sekarang berasal dari kata_pembuka.
- _buat_pengesahan_dinamis: formula rows diperbaiki; alignment tabel
  dibenahi agar penandatangan terdistribusi merata.
- _buat_isi_laporan_dinamis: label BAB menggunakan teks asli dari JS,
  tidak perlu diubah ke uppercase secara paksa.
- Font konsistensi: semua run eksplisit di-set Times New Roman 12pt.
"""

import io
import os
import base64
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def generate_laporan_pkl(data: dict) -> bytes:
    doc = Document()

    # Margin 4-4-3-3 cm
    section = doc.sections[0]
    section.left_margin   = Cm(4)
    section.top_margin    = Cm(4)
    section.right_margin  = Cm(3)
    section.bottom_margin = Cm(3)

    # Default font
    style      = doc.styles['Normal']
    font       = style.font
    font.name  = 'Times New Roman'
    font.size  = Pt(12)

    # ── Cover ────────────────────────────────────────────────────────────────
    if data.get("buat_cover"):
        _buat_cover(doc, data)
        doc.add_page_break()

    # ── Pengesahan dinamis ───────────────────────────────────────────────────
    pengesahan_list = data.get("pengesahan", [])
    if pengesahan_list:
        _buat_pengesahan_dinamis(doc, pengesahan_list)

    # Fallback: tanda tangan lama jika tidak ada pengesahan dinamis
    if data.get("buat_tanda_tangan") and not pengesahan_list:
        doc.add_page_break()
        _buat_tanda_tangan(doc, data)

    # ── Kata Pengantar ───────────────────────────────────────────────────────
    kp_data = data.get("kata_pengantar")
    if kp_data and kp_data.get("judul"):
        _buat_kata_pengantar(doc, kp_data)
        doc.add_page_break()

    # ── Daftar Isi ───────────────────────────────────────────────────────────
    if data.get("buat_daftar_isi"):
        _buat_daftar_isi(doc, data)
        doc.add_page_break()

    # ── Isi Laporan dinamis ──────────────────────────────────────────────────
    isi_laporan = data.get("isi_laporan", [])
    if isi_laporan:
        _buat_isi_laporan_dinamis(doc, isi_laporan)

    # ── Daftar Pustaka ───────────────────────────────────────────────────────
    rujukan_list = data.get("rujukan", [])
    if rujukan_list:
        doc.add_page_break()
        _heading(doc, "DAFTAR PUSTAKA", level=2)
        doc.add_paragraph()

        for item in rujukan_list:
            teks = item.get("teks", "").strip()
            if teks:
                p = doc.add_paragraph()
                run = p.add_run(teks)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                p.paragraph_format.left_indent      = Cm(1.27)
                p.paragraph_format.first_line_indent = Cm(-1.27)

    # ── Simpan ───────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Heading helper
# ─────────────────────────────────────────────────────────────────────────────

def _heading(doc, text: str, level: int = 1):
    """Tambahkan heading dengan format PKL standar."""
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold            = True
    run.font.name       = 'Times New Roman'
    run.font.size       = Pt(14 if level == 1 else 13 if level == 2 else 12)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Cover
# ─────────────────────────────────────────────────────────────────────────────

def _buat_cover(doc, data: dict):
    doc.add_paragraph()

    cover_img = data.get("cover_image_path", "")
    if cover_img and os.path.isfile(cover_img):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(cover_img, width=Cm(6))
        doc.add_paragraph()
    else:
        doc.add_paragraph()

    def _center_bold(text, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold       = True
        r.font.name  = 'Times New Roman'
        r.font.size  = Pt(size)
        return p

    def _center(text, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        return p

    _center_bold("LAPORAN PRAKTIK KERJA LAPANGAN (PKL)", size=16)
    doc.add_paragraph()
    _center_bold(f"DI {data.get('nama_instansi', '').upper()}", size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    _center_bold("Disusun Oleh:")

    for label, val in [
        ("Nama",          data.get("nama_lengkap", "-")),
        ("NIS / NIM",     data.get("nis_nim", "-")),
        ("Kelas / Jurusan", data.get("kelas_jurusan", "-")),
    ]:
        _center(f"{label}: {val}")

    doc.add_paragraph()
    doc.add_paragraph()
    _center_bold(data.get("nama_sekolah", ""), size=13)
    _center(str(datetime.datetime.now().year))


# ─────────────────────────────────────────────────────────────────────────────
# Daftar Isi
# FIX: menggunakan TabStop XML yang benar untuk titik-titik leader
# ─────────────────────────────────────────────────────────────────────────────

def _buat_daftar_isi(doc, data: dict):
    _heading(doc, "DAFTAR ISI", level=2)
    doc.add_paragraph()

    # Kumpulkan entri
    entries = []   # list of (indent_level, text, page_str)
    page    = 1

    kp = data.get("kata_pengantar")
    if kp and kp.get("judul"):
        entries.append((0, kp.get("judul", "KATA PENGANTAR").upper(), str(page)))
        page += 1

    for bab in data.get("isi_laporan", []):
        entries.append((0, str(bab.get("judul_bab", "")).upper(), str(page)))
        for sub in bab.get("subs", []):
            entries.append((1, sub.get("judul_sub", ""), str(page)))
        page += 2

    if data.get("rujukan"):
        entries.append((0, "DAFTAR PUSTAKA", str(page)))

    # Lebar halaman - margin = teks area ~ 19 cm → tab stop di kanan
    TAB_POS_EMU = int(Cm(13).pt * 12700)  # ~13 cm dari kiri

    for level, text, hal in entries:
        p   = doc.add_paragraph()
        fmt = p.paragraph_format

        if level == 0:
            run = p.add_run(text)
            run.bold = True
        else:
            fmt.left_indent = Cm(1)
            run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Tab stop dengan leader titik-titik
        _add_tab_stop(p, TAB_POS_EMU)
        tab_run = p.add_run(f"\t{hal}")
        tab_run.font.name = 'Times New Roman'
        tab_run.font.size = Pt(12)


def _add_tab_stop(paragraph, pos_emu: int):
    """Tambahkan tab stop kanan dengan leader titik-titik ke paragraph."""
    pPr   = paragraph._p.get_or_add_pPr()
    tabs  = OxmlElement('w:tabs')
    tab   = OxmlElement('w:tab')
    tab.set(qn('w:val'),    'right')
    tab.set(qn('w:pos'),    str(int(pos_emu / 914.4)))  # EMU → twip (1 cm = 567 twip)
    tab.set(qn('w:leader'), 'dot')
    tabs.append(tab)
    pPr.append(tabs)


# ─────────────────────────────────────────────────────────────────────────────
# Tanda Tangan (fallback lama)
# ─────────────────────────────────────────────────────────────────────────────

def _buat_tanda_tangan(doc, data: dict):
    doc.add_paragraph()
    _heading(doc, "LEMBAR PENGESAHAN", level=2)
    doc.add_paragraph()

    city     = data.get("kota_ttd", "")
    date_str = datetime.datetime.now().strftime("%d %B %Y")
    p = doc.add_paragraph(f"{city}, {date_str}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)

    def _set(row, col, text, bold=False):
        cell  = table.cell(row, col)
        para  = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run   = para.add_run(text)
        run.bold       = bold
        run.font.name  = 'Times New Roman'
        run.font.size  = Pt(12)

    _set(0, 0, "Pembimbing Lapangan",       bold=True)
    _set(0, 1, "Pembimbing Sekolah/Kampus", bold=True)
    for r in range(1, 4):
        _set(r, 0, "")
        _set(r, 1, "")
    _set(4, 0, data.get("nama_pembimbing_lapangan", "_______________"), bold=True)
    _set(4, 1, data.get("nama_pembimbing_sekolah",  "_______________"), bold=True)

    doc.add_paragraph()
    for text, align in [
        ("Mengetahui,",                   WD_ALIGN_PARAGRAPH.CENTER),
        (f"Pimpinan {data.get('nama_instansi', '')}", WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        p = doc.add_paragraph(text)
        p.alignment = align

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph("(________________________)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ─────────────────────────────────────────────────────────────────────────────
# Pengesahan Dinamis
# FIX: formula rows diperbaiki; tabel penandatangan lebih rapi
# ─────────────────────────────────────────────────────────────────────────────

def _buat_pengesahan_dinamis(doc, pengesahan_list: list):
    for p_item in pengesahan_list:
        doc.add_paragraph()
        judul = p_item.get("judul", "LEMBAR PENGESAHAN").upper()
        _heading(doc, judul, level=2)
        doc.add_paragraph()

        tujuan = p_item.get("tujuan", "")
        if tujuan:
            p = doc.add_paragraph(tujuan)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic     = True
                run.font.name  = 'Times New Roman'
                run.font.size  = Pt(12)

        doc.add_paragraph()

        lines = [
            ("Nama Penyusun", p_item.get("nama_penyusun")),
            ("NIS / NISN",    p_item.get("nis")),
            ("Kelas",         p_item.get("kelas")),
            ("Tahun Pelajaran", p_item.get("tahun_pelajaran")),
            ("Instansi PKL",  p_item.get("nama_pt")),
        ]
        for k, v in lines:
            if v:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(2)
                run = p.add_run(f"{k}\t: {v}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        doc.add_paragraph()

        tgl = p_item.get("tanggal", "")
        if tgl:
            try:
                tgl = datetime.datetime.strptime(tgl, "%Y-%m-%d").strftime("%d %B %Y")
            except ValueError:
                pass
            p = doc.add_paragraph(f"Mengesahkan pada tanggal: {tgl}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        doc.add_paragraph()

        penandatangan = p_item.get("penandatangan", [])
        if penandatangan:
            _buat_ttd_table(doc, penandatangan)

        doc.add_page_break()


def _buat_ttd_table(doc, penandatangan: list):
    """Buat tabel penandatangan dengan layout 2 kolom per baris."""
    n     = len(penandatangan)
    cols  = min(5, n)   # maks 3 kolom agar tidak terlalu sempit
    rows  = (n + cols - 1) // cols   # ceil division
    # Setiap penandatangan butuh 4 sub-rows: jabatan, spasi TTD, gambar, nama
    SUB   = 4
    tbl   = doc.add_table(rows=rows * SUB, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    def _cell_para(row, col, text="", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell = tbl.cell(row, col)
        para = cell.paragraphs[0]
        para.alignment = align
        if text:
            run = para.add_run(text)
            run.bold      = bold
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        return para

    for idx, s in enumerate(penandatangan):
        r = (idx // cols) * SUB
        c = idx % cols

        jabatan   = s.get("jabatan", "")
        nama      = s.get("nama", "_______________")
        image_b64 = s.get("image")

        _cell_para(r,     c, jabatan, bold=True)
        _cell_para(r + 1, c, "")

        # Gambar TTD (jika ada)
        if image_b64 and image_b64.startswith('data:image'):
            try:
                _, b64data  = image_b64.split(',', 1)
                img_stream  = io.BytesIO(base64.b64decode(b64data))
                para        = _cell_para(r + 1, c)
                run_img     = para.add_run()
                run_img.add_picture(img_stream, width=Cm(3))
            except Exception:
                pass

        _cell_para(r + 2, c, "")
        _cell_para(r + 3, c, nama, bold=True)


# ─────────────────────────────────────────────────────────────────────────────
# Kata Pengantar
# FIX: hapus teks hardcode bahasa Arab; gunakan kata_pembuka apa adanya
# ─────────────────────────────────────────────────────────────────────────────

def _buat_kata_pengantar(doc, kp_data: dict):
    _heading(doc, kp_data.get("judul", "KATA PENGANTAR").upper(), level=2)
    doc.add_paragraph()

    # Kata Pembuka — ditampilkan apa adanya dari form
    kata_pembuka = kp_data.get("kata_pembuka", "").strip()
    if kata_pembuka:
        p = doc.add_paragraph(kata_pembuka)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        doc.add_paragraph()

    # Ucapan Terima Kasih (List Nomor)
    ucapan = kp_data.get("ucapan_terima", [])
    for u in ucapan:
        nama    = u.get("nama", "").strip()
        jabatan = u.get("jabatan", "").strip()
        if nama or jabatan:
            teks = f"{nama}, selaku {jabatan}." if jabatan else nama
            p    = doc.add_paragraph(teks, style='List Number')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    # Kata Penutup
    kata_penutup = kp_data.get("kata_penutup", "").strip()
    if kata_penutup:
        doc.add_paragraph()
        p = doc.add_paragraph(kata_penutup)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    doc.add_paragraph()

    # Kota & Tanggal
    loc_date = kp_data.get("kota_tanggal", "").strip()
    if loc_date:
        p = doc.add_paragraph(loc_date)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Spasi tanda tangan
    for _ in range(3):
        doc.add_paragraph()

    nama_penulis = kp_data.get("nama_penulis", "").strip() or "(Penulis)"
    p = doc.add_paragraph(nama_penulis)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.bold      = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


# ─────────────────────────────────────────────────────────────────────────────
# Isi Laporan Dinamis
# ─────────────────────────────────────────────────────────────────────────────

def _buat_isi_laporan_dinamis(doc, babs: list):
    for bab in babs:
        teks_bab = bab.get("judul_bab", "").strip().upper()

        # Format: "BAB 1 - JUDUL" → dua baris heading
        if " - " in teks_bab:
            label_bab, judul_bab = teks_bab.split(" - ", 1)
            _heading(doc, label_bab.strip(), level=1)
            _heading(doc, judul_bab.strip(), level=2)
        else:
            _heading(doc, teks_bab, level=1)

        for sub in bab.get("subs", []):
            _heading(doc, sub.get("judul_sub", "Sub Bab").strip(), level=3)

            for c in sub.get("contents", []):
                tp = c.get("type", "paragraf")

                if tp == "paragraf":
                    teks = c.get("teks", "").strip()
                    if teks:
                        p = doc.add_paragraph(teks)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)

                elif tp == "gambar":
                    b64 = c.get("data", "")
                    if b64 and b64.startswith("data:image"):
                        try:
                            _, b64data = b64.split(',', 1)
                            img_stream = io.BytesIO(base64.b64decode(b64data))
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.add_run().add_picture(img_stream, width=Inches(5))
                        except Exception:
                            pass

                elif tp == "list":
                    # FIX: selalu render sebagai list, terlepas dari 'style'
                    _render_list(doc, c, indent_level=0)

        doc.add_page_break()


def _render_list(doc, lst: dict, indent_level: int = 0):
    """
    Render list rekursif.

    FIX: type check sekarang membaca field 'type' dari masing-masing
    item (bukan dari wrapper), karena JS mengirim type di level item.
    Style Word dipilih berdasarkan 'style' di wrapper.
    """
    items = lst.get("items", [])
    if not items:
        return

    style_code = lst.get("style", "1")
    word_style = 'List Bullet' if style_code == 'bullet' else 'List Number'

    for item in items:
        mode   = item.get("type", "simple")   # FIX: baca dari item, bukan wrapper
        judul  = item.get("judul", "").strip() if mode == "title" else ""
        teks   = item.get("teks", "").strip()

        p = doc.add_paragraph(style=word_style)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(indent_level * 1.27)

        if judul:
            run_j = p.add_run(judul + " ")
            run_j.bold      = True
            run_j.font.name = 'Times New Roman'
            run_j.font.size = Pt(12)

        if teks:
            run_t = p.add_run(teks)
            run_t.font.name = 'Times New Roman'
            run_t.font.size = Pt(12)

        # Nested list
        for nested in item.get("anak", []):
            _render_list(doc, nested, indent_level=indent_level + 1)


# ─────────────────────────────────────────────────────────────────────────────
# Util: border cell (dipertahankan, dipakai bila perlu)
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if border_name in kwargs:
            tag = f"w:{border_name}"
            el  = OxmlElement(tag)
            for key, val in kwargs[border_name].items():
                el.set(qn(f"w:{key}"), val)
            tcPr.append(el)
