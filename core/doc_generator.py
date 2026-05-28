import datetime
import io
import traceback

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .docx_utils import (
    add_tc_entry,
    coerce_dict_list,
    apply_all_docx_fixes,
    sanitize_inline_text,
    sanitize_paragraph_lines,
    set_run_fonts,
    take_first_non_empty,
)
from .image_engine import ImageEngine
from .layout_engine import LayoutEngine
from .page_engine import PageEngine
from .paragraph_engine import ParagraphEngine
from .signature_engine import SignatureEngine
from .styles import (
    FONT_NAME,
    TEXT_WIDTH_CM,
    apply_normal_style,
    configure_document_defaults,
    format_run,
)
from .toc_engine import TocEngine


def _sanitize_multiline_text(value):
    if value is None:
        return ""
    if isinstance(value, (list, tuple, set)):
        parts = [_sanitize_multiline_text(item) for item in value]
        return "\n\n".join(part for part in parts if part).strip()
    return sanitize_inline_text(value)


def _validate_data(data: dict):
    for field in ("nama_lengkap", "nama_instansi"):
        if not sanitize_inline_text(data.get(field)):
            return False, f"Field wajib '{field}' kosong atau tidak valid."
    return True, ""


def _bool(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "ya", "on"}


def _normalize_nested_list(item):
    if not isinstance(item, dict):
        return None
    nested_children = item.get("anak")
    if nested_children is None:
        nested_children = item.get("children", [])

    image_items = item.get("gambar_items") or []
    if item.get("imgId") or item.get("data") or item.get("imgData"):
        image_items = [
            *image_items,
            {
                "type": "gambar",
                "imgId": item.get("imgId", ""),
                "data": item.get("data") or item.get("imgData") or item.get("image") or item.get("src"),
                "caption": item.get("caption", ""),
            },
        ]

    item_type = sanitize_inline_text(item.get("type") or item.get("mode") or "simple").lower()
    if item_type == "complex":
        item_type = "title"

    return {
        "type": item_type,
        "judul": sanitize_inline_text(item.get("judul") or item.get("title") or ""),
        "teks": sanitize_inline_text(item.get("teks") or item.get("text") or ""),
        "gambar_items": [
            {
                "type": "gambar",
                "imgId": sanitize_inline_text(image.get("imgId", "")),
                "data": image.get("data") or image.get("imgData") or image.get("image") or image.get("src"),
                "caption": sanitize_inline_text(image.get("caption", "")),
            }
            for image in image_items
            if isinstance(image, dict) and (image.get("data") or sanitize_inline_text(image.get("imgId", "")))
        ],
        "anak": [
            nested_normalized
            for nested_normalized in (_normalize_list_wrapper(nested) for nested in nested_children)
            if nested_normalized
        ],
    }


def _normalize_list_wrapper(content):
    if not isinstance(content, dict):
        return None
    style = sanitize_inline_text(content.get("style", "1")) or "1"
    mode = sanitize_inline_text(content.get("mode", "")) or "simple"
    if mode == "complex":
        mode = "title"
    return {
        "type": "list",
        "style": style,
        "mode": mode,
        "items": [
            normalized_item
            for normalized_item in (_normalize_nested_list(item) for item in content.get("items", []))
            if normalized_item
        ],
    }


def _normalize_content(content):
    if not isinstance(content, dict):
        return None
    content_type = sanitize_inline_text(content.get("type") or "paragraf").lower()
    if content_type == "list":
        return _normalize_list_wrapper(content)
    normalized = {
        "type": content_type,
        "teks": _sanitize_multiline_text(content.get("teks", "")),
        "judul": sanitize_inline_text(content.get("judul", "")),
        "caption": sanitize_inline_text(content.get("caption", "")),
        "data": content.get("data") or content.get("image") or content.get("src"),
    }
    if content_type == "tabel":
        normalized["headers"] = [sanitize_inline_text(header) for header in content.get("headers", [])]
        normalized["rows"] = [
            [sanitize_inline_text(cell) for cell in row]
            for row in content.get("rows", [])
            if isinstance(row, list)
        ]
    return normalized


def _normalize_data(data: dict) -> dict:
    normalized = dict(data or {})
    simple_fields = [
        "nama_lengkap",
        "nis_nim",
        "kelas_jurusan",
        "nama_sekolah",
        "nama_instansi",
        "nama_pembimbing_lapangan",
        "nama_pembimbing_sekolah",
        "nama_kepala_sekolah",
        "kota_ttd",
        "tahun_ajaran",
        "cover_image_path",
        "cover_image_base64",
    ]
    for field in simple_fields:
        normalized[field] = sanitize_inline_text(normalized.get(field, ""))

    for field in ("buat_cover", "buat_daftar_isi", "buat_tanda_tangan"):
        normalized[field] = _bool(normalized.get(field))

    kata_pengantar = normalized.get("kata_pengantar", {})
    if not isinstance(kata_pengantar, dict):
        kata_pengantar = {}
    kata_pengantar["kata_pembuka"] = _sanitize_multiline_text(kata_pengantar.get("kata_pembuka", ""))
    kata_pengantar["kata_penutup"] = _sanitize_multiline_text(kata_pengantar.get("kata_penutup", ""))
    kata_pengantar["kota_tanggal"] = sanitize_inline_text(kata_pengantar.get("kota_tanggal", ""))
    kata_pengantar["nama_penulis"] = sanitize_inline_text(kata_pengantar.get("nama_penulis", ""))
    kata_pengantar["ucapan_terima"] = [
        {
            "nama": sanitize_inline_text(item.get("nama", "")),
            "jabatan": sanitize_inline_text(item.get("jabatan", "")),
        }
        for item in kata_pengantar.get("ucapan_terima", [])
        if isinstance(item, dict)
    ]
    judul_kp = sanitize_inline_text(kata_pengantar.get("judul", ""))
    has_kp_content = bool(
        kata_pengantar["kata_pembuka"]
        or kata_pengantar["kata_penutup"]
        or kata_pengantar["kota_tanggal"]
        or kata_pengantar["nama_penulis"]
        or kata_pengantar["ucapan_terima"]
    )
    kata_pengantar["judul"] = judul_kp or ("KATA PENGANTAR" if has_kp_content else "")
    normalized["kata_pengantar"] = kata_pengantar

    normalized["rujukan"] = [
        {"teks": _sanitize_multiline_text(item.get("teks", ""))}
        for item in normalized.get("rujukan", [])
        if isinstance(item, dict) and _sanitize_multiline_text(item.get("teks", ""))
    ]

    normalized["lampiran"] = [
        {
            "teks": _sanitize_multiline_text(item.get("teks", "")),
            "data": item.get("data") or item.get("image"),
        }
        for item in normalized.get("lampiran", [])
        if isinstance(item, dict)
    ]

    raw_pengesahan = coerce_dict_list(normalized.get("pengesahan"))
    normalized["pengesahan"] = []
    for item in raw_pengesahan:
        normalized["pengesahan"].append(
            {
                "judul": sanitize_inline_text(item.get("judul") or "LEMBAR PENGESAHAN"),
                "jenis_laporan": sanitize_inline_text(item.get("jenis_laporan") or "LAPORAN PRAKERIN"),
                "nama_pt": sanitize_inline_text(item.get("nama_pt") or normalized.get("nama_instansi")),
                "tujuan": sanitize_inline_text(item.get("tujuan", "")),
                "nama_penyusun": sanitize_inline_text(item.get("nama_penyusun") or normalized.get("nama_lengkap")),
                "nis": sanitize_inline_text(item.get("nis") or normalized.get("nis_nim")),
                "kelas": sanitize_inline_text(item.get("kelas") or normalized.get("kelas_jurusan")),
                "tahun_pelajaran": sanitize_inline_text(item.get("tahun_pelajaran") or normalized.get("tahun_ajaran")),
                "tanggal": sanitize_inline_text(item.get("tanggal", "")),
                "penandatangan": coerce_dict_list(item.get("penandatangan")),
            }
        )

    raw_isi_laporan = coerce_dict_list(normalized.get("isi_laporan"))
    normalized["isi_laporan"] = []
    for bab in raw_isi_laporan:
        normalized["isi_laporan"].append(
            {
                "judul_bab": sanitize_inline_text(bab.get("judul_bab", "")),
                "subs": [
                    {
                        "judul_sub": sanitize_inline_text(sub.get("judul_sub", "")),
                        "contents": [
                            normalized_content
                            for normalized_content in (_normalize_content(content) for content in sub.get("contents", []))
                            if normalized_content
                        ],
                    }
                    for sub in coerce_dict_list(bab.get("subs"))
                ],
            }
        )

    if not normalized["tahun_ajaran"]:
        year = datetime.datetime.now().year
        normalized["tahun_ajaran"] = f"{year}/{year + 1}"
    return normalized


def _split_bab_title(judul_bab: str, index: int):
    """
    Selalu kembalikan (label, judul) yang bersih.
    Output: ("BAB I", "PENDAHULUAN")

    Input yang ditangani:
      "PENDAHULUAN"            → ("BAB I", "PENDAHULUAN")
      "BAB I PENDAHULUAN"      → ("BAB I", "PENDAHULUAN")
      "BAB I - PENDAHULUAN"    → ("BAB I", "PENDAHULUAN")
      "BAB I — PENDAHULUAN"    → ("BAB I", "PENDAHULUAN")
      "BAB 1 PENDAHULUAN"      → ("BAB I", "PENDAHULUAN")
      "BAB 1 - PENDAHULUAN"    → ("BAB I", "PENDAHULUAN")
      "1 - PENDAHULUAN"        → ("BAB I", "PENDAHULUAN")
    """
    import re as _re

    ROMAN_LIST = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]

    def _to_roman(s):
        """Konversi angka Arab atau Romawi ke string Romawi standar."""
        s = s.strip().upper()
        if s.isdigit():
            i = int(s)
            return ROMAN_LIST[i - 1] if 1 <= i <= len(ROMAN_LIST) else s
        return s

    title = sanitize_inline_text(judul_bab).strip()
    upper = title.upper()

    # Normalkan semua tanda hubung (-, –, —) jadi spasi tunggal
    upper = _re.sub(r'\s*[-–—]\s*', ' ', upper).strip()

    # Pola 1: "BAB <angka_atau_romawi> <judul>"
    # \d+ HARUS dicoba sebelum pattern Romawi agar "1" tidak salah diparse
    m = _re.match(
        r'^BAB\s+(\d+|I{1,3}|IV|VI{0,3}|IX|XI{0,2}|XII)\s*(.*)',
        upper
    )
    if m:
        roman_str = _to_roman(m.group(1))
        body      = m.group(2).strip()
        return f"BAB {roman_str}", body

    # Pola 2: "<angka_atau_romawi> <judul>" (tanpa kata "BAB")
    m2 = _re.match(
        r'^(\d+|I{1,3}|IV|VI{0,3}|IX|XI{0,2}|XII)\s+(.*)',
        upper
    )
    if m2:
        roman_str = _to_roman(m2.group(1))
        body      = m2.group(2).strip()
        return f"BAB {roman_str}", body

    # Tidak ada prefix → auto-generate label dari index
    label = f"BAB {ROMAN_LIST[index - 1] if index - 1 < len(ROMAN_LIST) else index}"
    return label, upper


def _to_roman(number: int) -> str:
    pairs = (
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    result = []
    current = max(number, 1)
    for value, symbol in pairs:
        while current >= value:
            result.append(symbol)
            current -= value
    return "".join(result)


def _to_alpha(number: int, uppercase: bool = False) -> str:
    current = max(int(number), 1)
    chars = []
    while current:
        current -= 1
        chars.append(chr((65 if uppercase else 97) + (current % 26)))
        current //= 26
    return "".join(reversed(chars))


def _list_label(style: str, number: int) -> str:
    normalized_style = sanitize_inline_text(style) or "1"
    mapping = {
        "1": f"{number}.",
        "1.": f"{number}.",
        "1)": f"{number})",
        "(1)": f"({number})",
        "01.": f"{number:02d}.",
        "001.": f"{number:03d}.",
        "a": f"{_to_alpha(number)}.",
        "a.": f"{_to_alpha(number)}.",
        "a)": f"{_to_alpha(number)})",
        "(a)": f"({_to_alpha(number)})",
        "A": f"{_to_alpha(number, uppercase=True)}.",
        "A.": f"{_to_alpha(number, uppercase=True)}.",
        "A)": f"{_to_alpha(number, uppercase=True)})",
        "(A)": f"({_to_alpha(number, uppercase=True)})",
        "i": f"{_to_roman(number).lower()}.",
        "i.": f"{_to_roman(number).lower()}.",
        "i)": f"{_to_roman(number).lower()})",
        "(i)": f"({_to_roman(number).lower()})",
        "I": f"{_to_roman(number)}.",
        "I.": f"{_to_roman(number)}.",
        "I)": f"{_to_roman(number)})",
        "(I)": f"({_to_roman(number)})",
        "bullet": "•",
        "circle": "○",
        "square": "▪",
        "dash": "—",
        "arrow": "→",
        "none": "",
    }
    return mapping.get(normalized_style, f"{number}.")


def _strip_duplicate_prefix(text: str) -> str:
    value = sanitize_inline_text(text)
    value = value.replace("..", ".")
    parts = value.split()
    if len(parts) >= 2 and parts[0].rstrip(".").lower() == parts[1].rstrip(".").lower():
        return " ".join([parts[0]] + parts[2:])
    return value


def _ensure_alpha_prefix(text: str, index: int) -> str:
    value = _strip_duplicate_prefix(text)
    if not value:
        return f"{chr(64 + index)}."
    first = value.split()[0]
    if first.rstrip(".").isalpha() and len(first.rstrip(".")) == 1:
        return value
    return f"{chr(64 + index)}. {value}"


def _format_manual_city_date(value: str) -> str:
    text = sanitize_inline_text(value)
    if "," not in text:
        return text
    city, rest = text.split(",", 1)
    city = city.strip()
    rest = rest.strip()
    
    if rest and rest[0].isdigit():
        return f"{city}, {rest}"
    else:
        return f"{city},    {rest}" if rest else f"{city},    "


class ProfessionalPKLGenerator:
    def __init__(self, data: dict):
        self.data = data
        self.document = Document()
        configure_document_defaults(self.document)
        apply_all_docx_fixes(self.document)
        self.pages = PageEngine(self.document)
        self.paragraphs = ParagraphEngine(self.document)
        self.images = ImageEngine(self.document)
        self.layouts = LayoutEngine(self.document)
        self.toc = TocEngine(self.document)
        self.document._toc_engine = self.toc  # Inject for other engines
        self.signatures = SignatureEngine(self.document)

    def build(self) -> bytes:
        has_cover = self.data.get("buat_cover", True)
        has_preliminary = bool(self.data.get("pengesahan") or self.data.get("kata_pengantar", {}).get("judul") or self.data.get("buat_daftar_isi"))
        has_main = bool(self.data.get("isi_laporan") or self.data.get("rujukan") or self.data.get("lampiran"))

        current_section = self.document.sections[0]
        if has_cover:
            self.pages.setup_cover_section()
            self.build_cover()
        if has_preliminary:
            if has_cover:
                current_section = self.pages.add_section(WD_SECTION_START.NEW_PAGE)
            self.pages.setup_preliminary_section(current_section, start=1)
            self._build_preliminary_pages()
        if has_main:
            if has_cover or has_preliminary:
                current_section = self.pages.add_section(WD_SECTION_START.NEW_PAGE)
            self.pages.setup_body_section(current_section, start=1)
            self.build_body()
            self.build_references()
            self.build_lampiran()

        # Finalisasi TOC harus dilakukan setelah semua heading/section selesai dibuat.
        self.toc.finalize_toc(title="DAFTAR ISI")

        # Injeksi setting agar Word otomatis update fields saat dibuka.
        apply_all_docx_fixes(self.document)

        output = io.BytesIO()
        self.document.save(output)
        output.seek(0)
        return output.getvalue()

    def build_cover(self):
        title_specs = [
            ("LAPORAN", True, 20, 8),
            ("PRAKTIK KERJA LAPANGAN (PKL)", True, 18, 8),
            (f"DI {self.data.get('nama_instansi', '').upper()}", True, 18, 26),
        ]
        for text, bold, size, after in title_specs:
            paragraph = self.paragraphs.add_centered_line(text, bold=bold, size=Pt(size))
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(after)
            paragraph.paragraph_format.line_spacing = 1.5

        image_paragraph = self.images.insert_cover_image(
            take_first_non_empty([self.data.get("cover_image_path"), self.data.get("cover_image_base64")]),
            max_width_cm=7.5,
            max_height_cm=7.5,
        )
        if image_paragraph is not None:
            image_paragraph.paragraph_format.left_indent = Cm(0)
            image_paragraph.paragraph_format.first_line_indent = Cm(0)
            image_paragraph.paragraph_format.space_before = Pt(12)
            image_paragraph.paragraph_format.space_after = Pt(24)
            image_paragraph.paragraph_format.line_spacing = 1.0

        identity_specs = [
            ("Disusun Oleh", False, 14, 4),
            (self.data.get("nama_lengkap", ""), True, 15, 4),
            (self.data.get("nis_nim", ""), False, 13, 4),
            (self.data.get("kelas_jurusan", ""), False, 13, 18),
            (self.data.get("nama_sekolah", ""), True, 15, 4),
            (f"TAHUN PELAJARAN {self.data.get('tahun_ajaran', '')}", False, 14, 0),
        ]
        for text, bold, size, after in identity_specs:
            text = sanitize_inline_text(text)
            if not text:
                continue
            paragraph = self.paragraphs.add_centered_line(text, bold=bold, size=Pt(size))
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(after)
            paragraph.paragraph_format.line_spacing = 1.5

    def _build_preliminary_pages(self):
        approvals = self.data.get("pengesahan") or []
        chunks_built = 0
        for index, approval in enumerate(approvals):
            if chunks_built:
                self.pages.add_page_break()
            self.signatures.build_approval_page(approval, self.data)
            chunks_built += 1

        if self.data.get("kata_pengantar", {}).get("judul"):
            if chunks_built:
                self.pages.add_page_break()
            self.build_kata_pengantar()
            chunks_built += 1

        if self.data.get("buat_daftar_isi"):
            if chunks_built:
                self.pages.add_page_break()
            self.toc.insert_toc_placeholder()

    def build_kata_pengantar(self):
        kata_pengantar = self.data.get("kata_pengantar", {})
        self.paragraphs.add_section_heading(kata_pengantar.get("judul") or "KATA PENGANTAR")

        for block in sanitize_paragraph_lines(kata_pengantar.get("kata_pembuka", "")):
            paragraphs = self.paragraphs.add_body_text(block)
            for paragraph in paragraphs:
                paragraph.paragraph_format.space_after = Pt(4)

        for index, item in enumerate(kata_pengantar.get("ucapan_terima", []), start=1):
            name = sanitize_inline_text(item.get("nama"))
            role = sanitize_inline_text(item.get("jabatan"))
            if not (name or role):
                continue
            text = name if not role else f"{name}, selaku {role}."
            paragraph = self.paragraphs.add_list_item(f"{index}.", "", text, level=0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)

        for block in sanitize_paragraph_lines(kata_pengantar.get("kata_penutup", "")):
            paragraphs = self.paragraphs.add_body_text(block)
            for paragraph in paragraphs:
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)

        if kata_pengantar.get("kota_tanggal"):
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.space_before = Pt(36)
            paragraph.paragraph_format.space_after = Pt(0)
            format_run(paragraph.add_run(_format_manual_city_date(kata_pengantar["kota_tanggal"])), italic=False)

        if kata_pengantar.get("nama_penulis"):
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.space_before = Pt(48)
            paragraph.paragraph_format.space_after = Pt(0)
            format_run(paragraph.add_run(kata_pengantar["nama_penulis"]), bold=True, italic=False)

    def build_body(self):
        """Membangun isi laporan dengan Heading ter-unifikasi."""
        for bab_index, bab in enumerate(self.data.get("isi_laporan", []), start=1):
            if bab_index > 1:
                self.pages.add_page_break()

            label_bab, judul_bab = _split_bab_title(bab.get("judul_bab", ""), bab_index)

            # STANDAR PKL: "BAB I" baris pertama, "PENDAHULUAN" baris kedua
            # Keduanya CENTER, KAPITAL — dua paragraf Heading 1 terpisah
            # Spacing: space_after=0 di label agar rapat ke judul
            p_label = self.toc.insert_heading(label_bab.strip().upper(), level=1)
            p_label.paragraph_format.space_before = Pt(0)
            p_label.paragraph_format.space_after  = Pt(0)  # rapat ke judul

            if judul_bab.strip():
                p_judul = self.toc.insert_heading(judul_bab.strip().upper(), level=1)
                p_judul.paragraph_format.space_before = Pt(0)
                p_judul.paragraph_format.space_after  = Pt(18)  # jarak ke konten pertama

            # Sub bab (Heading 2)
            for sub_index, sub in enumerate(bab.get("subs", []), start=1):
                subheading = sanitize_inline_text(sub.get("judul_sub", ""))
                if subheading:
                    sub_text = _ensure_alpha_prefix(subheading, sub_index)
                    # PENTING: Gunakan Heading 2 asli Word
                    self.toc.insert_heading(sub_text, level=2)

                for content in sub.get("contents", []):
                    self._render_content(content, level=0)

    def _render_content(self, content: dict, level: int):
        content_type = content.get("type", "paragraf")
        if content_type == "paragraf":
            self.paragraphs.add_body_text(content.get("teks", ""))
            return

        # Heading/Subheading dari isi laporan juga wajib masuk TOC Word
        if content_type in {"heading", "subheading"}:
            heading_text = sanitize_inline_text(content.get("teks") or content.get("judul") or "")
            heading_level = int(content.get("level", 3) or 3)

            if heading_text:
                # Clamp ke 1..4, lalu map untuk Word Heading outline
                heading_level = max(1, min(4, heading_level))
                self.toc.insert_heading(heading_text, level=heading_level)
            return

        if content_type == "gambar":
            caption  = content.get("caption", "")
            img_data = (
                content.get("data") or
                content.get("imgData") or
                content.get("image") or
                content.get("src") or
                content.get("imgId")
            )
            self.images.insert_content_image(
                img_data,
                caption=caption,
                toc_entry_text=f"Gambar: {caption}" if caption else "Gambar",
            )
            return

        if content_type == "list":
            self._render_list(content, level=level)
            return

        if content_type == "tabel":
            self.layouts.add_data_table(content.get("headers", []), content.get("rows", []), content.get("caption", ""))

    def _render_list(self, lst: dict, level: int = 0):
        lst_mode = (lst.get("mode") or "simple").lower()

        for index, item in enumerate(lst.get("items", []), start=1):
            label = _list_label(lst.get("style", "1"), index)

            # Jika mode "title", maka judul_item wajib jadi Heading 3
            # agar masuk ke hierarki Daftar Isi secara profesional.
            item_mode = (item.get("type") or lst_mode or "simple").lower()
            title = item.get("judul", "") if item_mode in {"title", "complex"} else ""
            text = item.get("teks", "")

            if title:
                # PENTING: Judul item list murni dimasukkan sebagai Heading 3
                # agar otomatis masuk ke hierarki Daftar Isi secara profesional.
                self.toc.insert_heading(f"{label} {title}", level=3)
                
                # Buat paragraf teks (tanpa judul lagi karena sudah jadi Heading 3)
                if text:
                    self.paragraphs.add_body_text(text)
            else:
                # Item list biasa tanpa judul
                self.paragraphs.add_list_item(label, "", text, level=level)

            for image in item.get("gambar_items", []):
                img_data = (
                    image.get("data") or
                    image.get("imgData") or
                    image.get("image") or
                    image.get("src") or
                    image.get("imgId")
                )
                self.images.insert_content_image(
                    img_data,
                    caption=image.get("caption", ""),
                    max_width_cm=max(TEXT_WIDTH_CM - 2.5 - (level * 0.75), 7.0),
                    toc_entry_text=f"Gambar: {image.get('caption', '').strip()}" if image.get("caption", "").strip() else "Gambar",
                )

            for nested in item.get("anak", []):
                if isinstance(nested, dict):
                    self._render_list(nested, level=level + 1)


    def build_references(self):
        """Membangun Daftar Rujukan dengan hanging indent profesional."""
        references = self.data.get("rujukan", [])
        if not references:
            return
        
        self.pages.add_page_break()
        self.toc.insert_heading("DAFTAR RUJUKAN", level=1)

        for item in references:
            text = (item.get("teks") or "").strip()
            if not text:
                continue
                
            for block in sanitize_paragraph_lines(text):
                paragraph = self.document.add_paragraph()
                apply_normal_style(paragraph)
                
                # Setup Hanging Indent
                fmt = paragraph.paragraph_format
                fmt.left_indent = Cm(1.27)
                fmt.first_line_indent = Cm(-1.27)
                fmt.space_after = Pt(6)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                run = paragraph.add_run(block)
                set_run_fonts(run, font_name=FONT_NAME)

    def build_lampiran(self):
        """Membangun Lampiran."""
        lampiran = self.data.get("lampiran", [])
        if not lampiran:
            return
            
        self.pages.add_page_break()
        self.toc.insert_heading("LAMPIRAN", level=1)

        for item in lampiran:
            if item.get("judul"):
                self.toc.insert_heading(item["judul"], level=2)
            if item.get("teks"):
                self.paragraphs.add_body_text(item["teks"])
            if item.get("gambar"):
                self.images.insert_content_image(item["gambar"], caption=item.get("judul", ""), max_width_cm=TEXT_WIDTH_CM - 1.0)


def generate_laporan_pkl(data: dict) -> bytes:
    valid, message = _validate_data(data)
    if not valid:
        raise ValueError(message)

    normalized = _normalize_data(data)
    try:
        return ProfessionalPKLGenerator(normalized).build()
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"Gagal generate dokumen: {exc}\n{traceback.format_exc()}") from exc
